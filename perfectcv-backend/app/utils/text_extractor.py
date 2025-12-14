"""
Text extraction utilities for PDF and DOCX files.
Prioritizes PyMuPDF (fitz) for PDF and docx2txt for DOCX.
"""
import io
import logging
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# Import primary libraries
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) not available. Install with: pip install pymupdf")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logger.warning("docx2txt not available. Install with: pip install docx2txt")

# Fallback libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document
    DOCX_LIB_AVAILABLE = True
except ImportError:
    DOCX_LIB_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.info("OCR not available - install pytesseract, pillow, pdf2image for scanned PDF support")


class TextExtractor:
    """Extract text from PDF and DOCX files with prioritized clean extraction."""
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes, filename: str = "document.pdf") -> Tuple[str, str]:
        """
        Extract text from PDF using PyMuPDF (fitz) as primary.
        
        Args:
            file_bytes: PDF file content as bytes
            filename: Original filename for logging
            
        Returns:
            Tuple of (extracted_text, extraction_method)
        """
        text = ""
        method = "none"
        
        # 1. Primary: PyMuPDF (fitz)
        if FITZ_AVAILABLE:
            try:
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    text_parts = []
                    for page in doc:
                        # Extract text with layout preservation
                        page_text = page.get_text("text")
                        if page_text:
                            text_parts.append(page_text)
                    text = "\n\n".join(text_parts)
                    
                if text and len(text.strip()) > 50:
                    method = "pymupdf"
                    logger.info(f"Extracted {len(text)} chars from {filename} using PyMuPDF")
                    return text, method
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed for {filename}: {e}")

        # 2. Secondary: pdfplumber (Good for tables)
        if PDFPLUMBER_AVAILABLE:
            try:
                text_parts = []
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                text = "\n\n".join(text_parts)
                
                if text and len(text.strip()) > 50:
                    method = "pdfplumber"
                    logger.info(f"Extracted {len(text)} chars from {filename} using pdfplumber")
                    return text, method
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")

        # 3. Fallback: pdfminer.six
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract_text(io.BytesIO(file_bytes))
                if text and len(text.strip()) > 50:
                    method = "pdfminer"
                    logger.info(f"Extracted {len(text)} chars from {filename} using pdfminer")
                    return text, method
            except Exception as e:
                logger.warning(f"pdfminer extraction failed: {e}")

        # 4. Last Resort: OCR
        if OCR_AVAILABLE:
            try:
                text = TextExtractor._extract_with_ocr(file_bytes)
                if text and len(text.strip()) > 50:
                    method = "ocr"
                    logger.info(f"Extracted {len(text)} chars from {filename} using OCR")
                    return text, method
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")

        if not text:
            logger.error(f"Failed to extract text from {filename} using all available methods")
            raise ValueError("Could not extract text from PDF. File may be corrupted or scanned without OCR support.")
        
        return text, method

    @staticmethod
    def _extract_with_ocr(file_bytes: bytes) -> str:
        """Extract text using OCR (for scanned PDFs)."""
        import tempfile
        import os
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(tmp_path)
            text_parts = []
            
            for i, image in enumerate(images):
                logger.info(f"Performing OCR on page {i+1}/{len(images)}")
                page_text = pytesseract.image_to_string(image)
                if page_text:
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    @staticmethod
    def extract_from_docx(file_bytes: bytes, filename: str = "document.docx") -> Tuple[str, str]:
        """
        Extract text from DOCX using docx2txt as primary.
        
        Args:
            file_bytes: DOCX file content as bytes
            filename: Original filename for logging
            
        Returns:
            Tuple of (extracted_text, extraction_method)
        """
        text = ""
        method = "none"

        # 1. Primary: docx2txt
        if DOCX2TXT_AVAILABLE:
            try:
                # docx2txt process expects a file path or file-like object, but usually works best with temporary file
                # However, the library 'docx2txt' typically exposes process(filename, html=False)
                # We can try to use a temporary file to be safe
                import tempfile
                import os
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_file.write(file_bytes)
                    tmp_path = tmp_file.name
                
                try:
                    text = docx2txt.process(tmp_path)
                    if text and len(text.strip()) > 20:
                        method = "docx2txt"
                        logger.info(f"Extracted {len(text)} chars from {filename} using docx2txt")
                        return text, method
                finally:
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass
            except Exception as e:
                logger.warning(f"docx2txt extraction failed: {e}")

        # 2. Fallback: python-docx
        if DOCX_LIB_AVAILABLE:
            try:
                doc = Document(io.BytesIO(file_bytes))
                text_parts = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text_parts.append(row_text)
                
                text = "\n".join(text_parts)
                if text and len(text.strip()) > 20:
                    method = "python-docx"
                    logger.info(f"Extracted {len(text)} chars from {filename} using python-docx")
                    return text, method
            except Exception as e:
                logger.warning(f"python-docx extraction failed: {e}")

        if not text:
            raise ValueError(f"Could not extract text from DOCX: {filename}")
            
        return text, method

    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """
        Auto-detect file type and extract text.
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, extraction_method)
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return TextExtractor.extract_from_pdf(file_bytes, filename)
        elif filename_lower.endswith(('.docx', '.doc')):
            return TextExtractor.extract_from_docx(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported file type: {filename}. Supported: PDF, DOCX")
