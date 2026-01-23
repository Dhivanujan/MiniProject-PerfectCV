import io
import os
import re
import json
import logging
import zipfile
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET
from pdfminer.high_level import extract_text as pdf_extract_text
from app.utils import cleaner as _cleaner
from pdfminer.layout import LAParams
from docx import Document
import phonenumbers
from fpdf import FPDF
import google.generativeai as genai
from app.utils.ai_utils import setup_gemini, get_valid_model, improve_sentence, get_generative_model, generate_with_retry
from app.utils.nlp_utils import load_spacy_model, extract_entities, classify_header_nlp

# PDF extraction - use PyMuPDF
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("PyMuPDF not available. Install: pip install pymupdf")

logger = logging.getLogger(__name__)


ALLOWED_EXTENSIONS = {"pdf", "doc", "docx"}


def needs_ai_extraction(contact_info: dict) -> bool:
    """Check if AI extraction is needed for incomplete contact info.
    
    Args:
        contact_info: Dictionary with name, email, phone fields
        
    Returns:
        True if any critical field (name, email, phone) is missing
    """
    critical_fields = ["name", "email", "phone"]
    missing_fields = [f for f in critical_fields if not contact_info.get(f)]
    
    if missing_fields:
        logger.warning(f"⚠ Incomplete extraction. Missing: {', '.join(missing_fields)}")
        return True
    
    logger.info(f"✓ Complete extraction: name={bool(contact_info.get('name'))}, "
                f"email={bool(contact_info.get('email'))}, phone={bool(contact_info.get('phone'))}")
    return False


def validate_contact_info(contact_info: dict) -> dict:
    """Validate and score contact information completeness.
    
    Args:
        contact_info: Dictionary with contact fields
        
    Returns:
        Dictionary with validation results and completeness score
    """
    import re
    
    result = {
        'valid_name': False,
        'valid_email': False,
        'valid_phone': False,
        'completeness': 0,
        'score': 0
    }
    
    # Validate name (at least 2 words, 3+ characters each)
    name = contact_info.get('name', '').strip()
    if name and len(name) >= 6:
        words = name.split()
        if len(words) >= 2 and all(len(w) >= 2 for w in words):
            result['valid_name'] = True
    
    # Validate email (proper format)
    email = contact_info.get('email', '').strip()
    if email and '@' in email and '.' in email.split('@')[-1]:
        if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
            result['valid_email'] = True
    
    # Validate phone (has digits, reasonable length)
    phone = contact_info.get('phone', '').strip()
    if phone:
        digits = re.sub(r'\D', '', phone)
        if 7 <= len(digits) <= 15:  # Valid phone number length
            result['valid_phone'] = True
    
    # Calculate completeness
    valid_count = sum([result['valid_name'], result['valid_email'], result['valid_phone']])
    result['completeness'] = valid_count
    result['score'] = (valid_count / 3.0) * 100
    
    return result


def extract_name_with_spacy(text: str) -> str:
    """Extract name using spaCy PERSON entity recognition ONLY.
    
    NO REGEX ALLOWED for name extraction.
    
    Args:
        text: CV text (searches first 500 characters only)
        
    Returns:
        Extracted name or empty string
    """
    from app.utils.nlp_utils import get_nlp
    
    nlp = get_nlp()
    if not nlp:
        logger.warning("⚠ spaCy not available, cannot extract name")
        return ""
    
    # Blacklist of common tech terms that spaCy might misidentify as names
    TECH_BLACKLIST = {
        'spring boot', 'react', 'angular', 'vue', 'node', 'nodejs', 'java', 
        'python', 'javascript', 'typescript', 'spring', 'django', 'flask',
        'docker', 'kubernetes', 'aws', 'azure', 'mongodb', 'mysql', 'postgresql',
        'redis', 'kafka', 'jenkins', 'github', 'gitlab', 'jira', 'confluence',
        'tensorflow', 'pytorch', 'keras', 'pandas', 'numpy', 'scikit', 'opencv',
        'express', 'fastapi', 'laravel', 'symfony', 'rails', 'ruby', 'php',
        'c++', 'c#', 'golang', 'rust', 'kotlin', 'swift', 'objective-c',
        'android', 'ios', 'linux', 'windows', 'macos', 'ubuntu', 'centos'
    }
    
    # Search only first 500 characters (name is usually at top)
    search_text = text[:500] if len(text) > 500 else text
    
    try:
        doc = nlp(search_text)
        
        # Find PERSON entities - collect all candidates
        candidates = []
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                name_lower = name.lower()
                words = name.split()
                
                # Check blacklist - skip technology/framework names
                if name_lower in TECH_BLACKLIST:
                    logger.debug(f"Skipping tech term identified as person: '{name}'")
                    continue
                
                # Skip common non-name terms
                if name_lower in ['name', 'resume', 'cv', 'curriculum vitae']:
                    continue
                
                # Less restrictive: Accept single names OR multi-word names
                # Just require reasonable character count
                if len(name) >= 3:
                    # Skip if it's mostly numbers or special characters
                    if sum(c.isalpha() for c in name) / len(name) > 0.5:
                        # Additional validation: Check if it contains common name patterns
                        # Real names typically have proper capitalization and reasonable word count
                        is_valid_name = True
                        
                        # Skip if single word and all caps (likely acronym or section header)
                        if len(words) == 1 and name.isupper() and len(name) > 4:
                            is_valid_name = False
                            logger.debug(f"Skipping all-caps single word: '{name}'")
                        
                        if is_valid_name:
                            candidates.append((name, len(words), ent.start_char))
                            logger.debug(f"Name candidate: '{name}' ({len(words)} words, pos {ent.start_char})")
        
        if candidates:
            # Prefer multi-word names that appear early in document
            # Sort by: (1) word count DESC (prefer full names), (2) position ASC (prefer top)
            candidates.sort(key=lambda x: (-x[1], x[2]))
            best_name = candidates[0][0]
            logger.info(f"✓ Name extracted via spaCy: {best_name} (from {len(candidates)} candidates)")
            return best_name
        
        logger.warning("⚠ No valid PERSON entity found in first 500 chars")
        return ""
    except Exception as e:
        logger.warning(f"⚠ spaCy name extraction failed: {e}")
        return ""


def extract_email_with_regex(text: str) -> str:
    """Extract email using regex after text normalization."""
    import re
    
    # Standard email pattern
    pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
    matches = re.findall(pattern, text)
    
    if matches:
        # Return first valid-looking email
        for email in matches:
            if not email.startswith('.') and not email.endswith('.'):
                logger.info(f"✓ Email extracted: {email}")
                return email
    
    return ""


def extract_phone_with_regex(text: str) -> str:
    """Extract phone using regex after text normalization.
    
    Supports: +country codes, (), spaces, dashes
    Example: +94 7720272019, (555) 123-4567, +1-555-123-4567
    """
    import re
    
    # Comprehensive phone pattern
    # Matches: +1234567890, +1 234 567 8900, (555) 123-4567, 555-123-4567, etc.
    pattern = r'\+?\d{1,4}[\s.-]?\(?\d{1,4}\)?[\s.-]?\d{1,4}[\s.-]?\d{1,4}[\s.-]?\d{0,4}'
    matches = re.findall(pattern, text)
    
    if matches:
        # Find the longest match (likely most complete)
        best_match = max(matches, key=lambda x: len(re.sub(r'\D', '', x)))
        digits = re.sub(r'\D', '', best_match)
        
        # Validate phone number length (7-15 digits)
        if 7 <= len(digits) <= 15:
            phone = best_match.strip()
            logger.info(f"✓ Phone extracted: {phone}")
            return phone
    
    return ""


def extract_contact_info_basic(text: str) -> dict:
    """Extract basic contact info using spaCy (name) and regex (email/phone).
    
    Args:
        text: Normalized CV text
        
    Returns:
        Dictionary with name, email, phone, linkedin, github
    """
    contact_info = {
        'name': '',
        'email': '',
        'phone': '',
        'linkedin': '',
        'github': ''
    }
    
    # Extract name using spaCy ONLY (NO REGEX)
    contact_info['name'] = extract_name_with_spacy(text)
    
    # Extract email and phone using regex
    contact_info['email'] = extract_email_with_regex(text)
    contact_info['phone'] = extract_phone_with_regex(text)
    
    # Extract LinkedIn
    import re
    linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+/?'
    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_match:
        contact_info['linkedin'] = linkedin_match.group(0)
    
    # Extract GitHub
    github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+/?'
    github_match = re.search(github_pattern, text, re.IGNORECASE)
    if github_match:
        contact_info['github'] = github_match.group(0)
    
    return contact_info


def format_extracted_text_with_sections(raw_text: str) -> dict:
    """Format extracted CV text with clear section headers for better readability.
    
    Args:
        raw_text: Raw extracted CV text
        
    Returns:
        Dictionary with:
        - formatted_text: Text with clear section headers
        - sections: Dictionary of categorized sections
        - section_order: List of section names in order
    """
    from app.utils.cv_utils import build_standardized_sections
    
    # Build structured sections
    structured = build_standardized_sections(raw_text)
    
    # Format with clear headers
    formatted_parts = []
    section_map = {}
    section_order = []
    
    # Personal Information / Contact
    contact = structured.get('contact_information', {})
    if contact and any(contact.values()):
        section_text = "═══════════════════════════════════════════════════\n"
        section_text += "📧 PERSONAL INFORMATION\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        
        if contact.get('name'):
            section_text += f"Name: {contact['name']}\n"
        if contact.get('email'):
            section_text += f"Email: {contact['email']}\n"
        if contact.get('phone'):
            section_text += f"Phone: {contact['phone']}\n"
        if contact.get('linkedin'):
            section_text += f"LinkedIn: {contact['linkedin']}\n"
        if contact.get('github'):
            section_text += f"GitHub: {contact['github']}\n"
        if contact.get('address'):
            section_text += f"Address: {contact['address']}\n"
        
        formatted_parts.append(section_text)
        section_map['Personal Information'] = section_text
        section_order.append('Personal Information')
    
    # Professional Summary
    summary = structured.get('professional_summary', '').strip()
    if summary:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "💼 PROFESSIONAL SUMMARY\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        section_text += summary + "\n"
        
        formatted_parts.append(section_text)
        section_map['Professional Summary'] = section_text
        section_order.append('Professional Summary')
    
    # Skills
    skills = structured.get('skills', {})
    if skills and (skills.get('technical') or skills.get('soft') or skills.get('languages_skills')):
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "🔧 SKILLS\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        
        if skills.get('technical'):
            section_text += "Technical Skills:\n"
            section_text += "  • " + "\n  • ".join(skills['technical']) + "\n\n"
        
        if skills.get('soft'):
            section_text += "Soft Skills:\n"
            section_text += "  • " + "\n  • ".join(skills['soft']) + "\n\n"
        
        if skills.get('languages_skills'):
            section_text += "Programming Languages:\n"
            section_text += "  • " + "\n  • ".join(skills['languages_skills']) + "\n"
        
        formatted_parts.append(section_text)
        section_map['Skills'] = section_text
        section_order.append('Skills')
    
    # Work Experience
    experience = structured.get('work_experience', [])
    if experience:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "💼 WORK EXPERIENCE\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        
        for exp in experience:
            if isinstance(exp, dict):
                section_text += f"• {exp.get('role', 'Position')}"
                if exp.get('company'):
                    section_text += f" at {exp['company']}"
                if exp.get('years'):
                    section_text += f" ({exp['years']})"
                section_text += "\n"
                
                if exp.get('description'):
                    section_text += f"  {exp['description']}\n"
                section_text += "\n"
        
        formatted_parts.append(section_text)
        section_map['Work Experience'] = section_text
        section_order.append('Work Experience')
    
    # Projects
    projects = structured.get('projects', [])
    if projects:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "🚀 PROJECTS\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        
        for proj in projects:
            if isinstance(proj, dict):
                section_text += f"• {proj.get('name', 'Project')}\n"
                if proj.get('description'):
                    section_text += f"  {proj['description']}\n"
                if proj.get('technologies'):
                    section_text += f"  Technologies: {', '.join(proj['technologies'])}\n"
                section_text += "\n"
        
        formatted_parts.append(section_text)
        section_map['Projects'] = section_text
        section_order.append('Projects')
    
    # Education
    education = structured.get('education', [])
    if education:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "🎓 EDUCATION\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        
        for edu in education:
            if isinstance(edu, dict):
                section_text += f"• {edu.get('degree', 'Degree')}"
                if edu.get('institution'):
                    section_text += f" - {edu['institution']}"
                if edu.get('year'):
                    section_text += f" ({edu['year']})"
                section_text += "\n"
                
                if edu.get('details'):
                    section_text += f"  {edu['details']}\n"
                section_text += "\n"
        
        formatted_parts.append(section_text)
        section_map['Education'] = section_text
        section_order.append('Education')
    
    # Certifications
    certifications = structured.get('certifications', [])
    if certifications:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "📜 CERTIFICATIONS\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        section_text += "  • " + "\n  • ".join(certifications) + "\n"
        
        formatted_parts.append(section_text)
        section_map['Certifications'] = section_text
        section_order.append('Certifications')
    
    # Achievements
    achievements = structured.get('achievements', [])
    if achievements:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "🏆 ACHIEVEMENTS\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        section_text += "  • " + "\n  • ".join(achievements) + "\n"
        
        formatted_parts.append(section_text)
        section_map['Achievements'] = section_text
        section_order.append('Achievements')
    
    # Languages
    languages = structured.get('languages', [])
    if languages:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "🌍 LANGUAGES\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        section_text += "  • " + "\n  • ".join(languages) + "\n"
        
        formatted_parts.append(section_text)
        section_map['Languages'] = section_text
        section_order.append('Languages')
    
    # Additional Information
    additional = structured.get('additional_information', '').strip()
    if additional:
        section_text = "\n═══════════════════════════════════════════════════\n"
        section_text += "ℹ️  ADDITIONAL INFORMATION\n"
        section_text += "═══════════════════════════════════════════════════\n\n"
        section_text += additional + "\n"
        
        formatted_parts.append(section_text)
        section_map['Additional Information'] = section_text
        section_order.append('Additional Information')
    
    # Combine all sections
    formatted_text = "\n".join(formatted_parts)
    
    return {
        'formatted_text': formatted_text,
        'sections': section_map,
        'section_order': section_order,
        'structured_data': structured
    }


def allowed_file(filename):
    """Return True if filename has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


DOMAIN_KEYWORDS = {
    "software": [
        "python", "java", "c++", "c#", "javascript", "react", "node", "docker", "kubernetes",
        "aws", "azure", "gcp", "ci/cd", "rest", "graphql", "microservices", "sql", "nosql"
    ],
    "data_science": [
        "python", "r", "sql", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "machine learning",
        "nlp", "computer vision", "statistics", "data visualization"
    ],
    "product": ["roadmap", "roadmapping", "stakeholders", "a/b testing", "user research", "metrics", "okrs", "prds"],
    "design": ["ux", "ui", "figma", "prototyping", "wireframes", "user research", "visual design"],
    "marketing": ["seo", "content", "campaign", "analytics", "growth", "google ads", "facebook ads", "email marketing"]
}

STANDARD_SECTION_ORDER: List[Tuple[str, str]] = [
    ("contact_information", "Personal Information"),
    ("professional_summary", "Professional Summary"),
    ("skills", "Skills"),
    ("work_experience", "Work Experience"),
    ("projects", "Projects"),
    ("education", "Education"),
    ("certifications", "Certifications"),
    ("achievements", "Achievements / Awards"),
    ("languages", "Languages"),
    ("volunteer_experience", "Volunteer Experience"),
    ("additional_information", "Additional Information"),
]

SOFT_SKILL_KEYWORDS: Tuple[str, ...] = (
    "communication", "leadership", "collaboration", "team", "problem", "critical", "creative",
    "adaptability", "negotiation", "stakeholder", "mentoring", "mentorship", "coaching", "presentation",
    "interpersonal", "time management", "conflict", "organization", "organisational", "empathy",
)

TECH_SKILL_HINTS: Tuple[str, ...] = (
    "python", "java", "c++", "c#", "go", "javascript", "typescript", "sql", "nosql", "aws", "azure",
    "gcp", "docker", "kubernetes", "linux", "react", "angular", "vue", "node", "django", "flask",
    "spring", "ml", "machine", "data", "analytics", "cloud", "devops", "git", "spark", "hadoop",
    "tensorflow", "pytorch", "api", "rest", "graphql", "microservice", "etl", "pandas", "numpy",
    "tableau", "power bi", "excel", "jira", "confluence", "salesforce", "php", "html", "css",
)

LANGUAGE_NAMES: Tuple[str, ...] = (
    "english", "french", "spanish", "german", "hindi", "mandarin", "cantonese", "arabic", "portuguese",
    "italian", "japanese", "korean", "tamil", "telugu", "malayalam", "urdu", "bengali", "marathi",
    "russian", "polish", "dutch", "swedish", "danish", "norwegian", "turkish", "thai", "vietnamese",
    "indonesian", "malay", "finnish", "greek", "hebrew", "punjabi", "romanian", "czech", "slovak",
)

SECTION_SYNONYMS: Dict[str, Tuple[str, ...]] = {
    "about": (
        "about",
        "summary",
        "professional summary",
        "profile",
        "career summary",
        "objective",
        "career objective",
        "about me",
        "contact",
        "contact information",
        "contact details",
        "personal information",
        "highlights",
    ),
    "skills": (
        "skills",
        "technical skills",
        "key skills",
        "core skills",
        "competencies",
        "skill set",
        "areas of expertise",
        "summary of skills",
    ),
    "experience": (
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "career history",
        "employment",
        "work history",
        "relevant experience",
        "experience highlights",
    ),
    "projects": ("projects", "project experience", "project work", "portfolio", "case studies"),
    "education": (
        "education",
        "academic background",
        "qualifications",
        "education & training",
        "academic history",
        "academics",
    ),
    "achievements": (
        "achievements",
        "accomplishments",
        "awards",
        "honors",
        "recognitions",
        "events",
        "extracurricular",
        "activities",
    ),
    "certifications": (
        "certifications",
        "certification",
        "licenses",
        "licences",
        "credentials",
        "certificates",
        "training",
        "licensure",
    ),
    "volunteer": (
        "volunteer",
        "volunteer experience",
        "volunteer work",
        "community service",
        "community involvement",
        "service",
    ),
    "languages": (
        "languages",
        "language proficiency",
        "language skills",
        "spoken languages",
        "languages known",
    ),
    "other": (
        "additional information",
        "other",
        "interests",
        "activities",
        "extra",
        "hobbies",
        "additional details",
    ),
}

ALL_SECTION_KEYWORDS = tuple(sorted({kw for values in SECTION_SYNONYMS.values() for kw in values}))
ALL_SECTION_HEADINGS_PATTERN = "|".join(re.escape(keyword) for keyword in ALL_SECTION_KEYWORDS)

SECTION_KEYS: Tuple[str, ...] = (
    "about",
    "summary",
    "skills",
    "experience",
    "education",
    "projects",
    "achievements",
    "certifications",
    "volunteer",
    "languages",
    "other",
)

INLINE_HEADING_PATTERN = re.compile(r"^(?P<label>[A-Za-z][\w &+/().']{1,80})\s*[:\-–]\s*(?P<body>.+)$")
BULLET_PREFIXES: Tuple[str, ...] = ("-", "*", "•")

MONTH_PATTERN = r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
DOB_DATE_PATTERN = re.compile(
    rf"(\b\d{{1,2}}[-/.]\d{{1,2}}[-/.](?:19|20)\d{{2}}\b|\b(?:19|20)\d{{2}}[-/.]\d{{1,2}}[-/.]\d{{1,2}}\b|\b\d{{1,2}}\s+{MONTH_PATTERN}\s+(?:19|20)\d{{2}}\b|\b{MONTH_PATTERN}\s+\d{{1,2}},?\s+(?:19|20)\d{{2}}\b)",
    re.IGNORECASE,
)
DOB_LABEL_PATTERN = re.compile(r"\b(?:dob|d\.o\.b|date of birth|birthdate|birthday|born)\b", re.IGNORECASE)
ADDRESS_KEYWORDS: Tuple[str, ...] = (
    "address",
    "resides",
    "residence",
    "residing",
    "living at",
    "living in",
    "permanent address",
    "current address",
    "mailing address",
    "postal address",
    "home address",
)


def _normalize_heading_label(label: str) -> str:
    cleaned = (label or "").strip().lower()
    cleaned = re.sub(r"[:\-–]+$", "", cleaned)
    cleaned = re.sub(r"[^a-z0-9&+/ ]+", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _build_section_heading_index() -> Dict[str, str]:
    index: Dict[str, str] = {}
    for key, keywords in SECTION_SYNONYMS.items():
        for keyword in keywords:
            normalized = _normalize_heading_label(keyword)
            if normalized and normalized not in index:
                index[normalized] = key
    return index


SECTION_HEADING_INDEX = _build_section_heading_index()


def _heading_lookup(label: str) -> Optional[str]:
    normalized = _normalize_heading_label(label)
    if not normalized:
        return None
    if normalized in SECTION_HEADING_INDEX:
        return SECTION_HEADING_INDEX[normalized]
    for keyword, key in SECTION_HEADING_INDEX.items():
        if normalized.startswith(keyword) and len(normalized.split()) <= len(keyword.split()) + 2:
            return key
    return None


def _extract_dob_from_text(text: str, allow_year_only: bool = False) -> str:
    if not text:
        return ""
    match = DOB_DATE_PATTERN.search(text)
    if match:
        return match.group(0).strip(" .,;|-")
    if allow_year_only:
        year_match = re.search(r"(19|20)\d{2}", text)
        if year_match:
            return year_match.group(0)
    return ""


def normalize_text(text):
    """Clean and normalize text while preserving meaningful spacing and structure."""
    if not text:
        return ""
        
    # Replace multiple spaces with single space
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Fix common PDF extraction artifacts
    text = text.replace('\x00', '') # Null bytes
    text = text.replace('\f', '\n') # Form feeds
    
    # Preserve newlines that likely indicate sections or list items
    text = re.sub(r'([.!?])\s*\n\s*([A-Z])', r'\1\n\n\2', text)
    
    # Ensure list items and bullets start on new lines
    text = re.sub(r'\s*([•\-\*]|\d+\.)\s*', r'\n\1 ', text)
    
    # Fix collapsed words (missing spaces after punctuation)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)
    
    # Fix email addresses that may be split
    text = re.sub(r'(\w+)\s+@\s+(\w+)', r'\1@\2', text)
    
    # Remove repeated newlines while preserving paragraph breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Clean up extra spaces in lines
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def _safe_decode_bytes(data: bytes, encoding: str = "utf-8") -> str:
    """Best-effort decode for binary CV uploads."""
    try:
        return data.decode(encoding, errors="ignore")
    except Exception:
        try:
            return data.decode("latin-1", errors="ignore")
        except Exception:
            return ""


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return normalize_text('\n'.join(full_text))
    except Exception as exc:
        logger.warning("DOCX extraction failed with python-docx: %s", exc)
        return normalize_text(_safe_decode_bytes(file_bytes))


def extract_text_from_doc_bytes(file_bytes: bytes) -> str:
    """Fallback extraction for legacy DOC files using best-effort decoding."""
    # Binary .doc format is proprietary; without external tools we attempt best-effort decoding.
    decoded = _safe_decode_bytes(file_bytes)
    return normalize_text(decoded)


def _dedupe_preserve_order(items: Iterable[str]) -> List[str]:
    """Return a list with duplicates removed while preserving original ordering."""
    seen = set()
    ordered: List[str] = []
    for item in items:
        if not item:
            continue
        normalized = item.strip()
        if not normalized:
            continue
        key = normalized.lower()
        if key in seen:
            continue
        seen.add(key)
        ordered.append(normalized)
    return ordered


def _split_section_lines(section_text: str) -> List[str]:
    """Convert multiline/bulleted section text into normalized line items."""
    if not section_text:
        return []
    entries: List[str] = []
    for raw_line in section_text.split("\n"):
        cleaned = raw_line.strip().lstrip("-•* ").strip()
        if cleaned:
            entries.append(cleaned)
    return entries

def _extract_section_by_keywords(text: str, keywords: Sequence[str]) -> str:
    """Best-effort extraction of a section block given keyword variants."""
    if not keywords:
        return ""
    heading_group = "|".join(re.escape(keyword) for keyword in keywords if keyword)
    if not heading_group:
        return ""

    inline_pattern = rf"(?:^|\n)\s*(?:{heading_group})\b[^\n]*[:\-]\s*(?P<inline>[^\n\r]+)"
    match = re.search(inline_pattern, text, flags=re.IGNORECASE)
    if match:
        return match.group("inline").strip()

    stop_pattern = rf"(?=\n\s*(?:{ALL_SECTION_HEADINGS_PATTERN})\b|$)" if ALL_SECTION_HEADINGS_PATTERN else r"(?=$)"
    block_pattern = rf"(?:^|\n)\s*(?:{heading_group})\b[^\n]*\n(?P<body>.*?){stop_pattern}"
    match = re.search(block_pattern, text, flags=re.IGNORECASE | re.DOTALL)
    if match:
        return match.group("body").strip()
    return ""


def _augment_sections_from_keywords(text: str, sections: Dict[str, str]) -> None:
    """Populate missing sections using keyword heuristics."""
    for key, keywords in SECTION_SYNONYMS.items():
        if not keywords or sections.get(key):
            continue
        snippet = _extract_section_by_keywords(text, keywords)
        if snippet:
            sections[key] = (sections.get(key, "") + "\n" + snippet).strip()


def _infer_skills_from_text(text: str) -> List[str]:
    """Collect technical keywords within free-form text as fallback skills."""
    found: List[str] = []
    seen = set()
    text_lower = text.lower()
    for keyword in TECH_SKILL_HINTS:
        if keyword and keyword.lower() in text_lower:
            cleaned = keyword.strip()
            normalized = cleaned if cleaned.isupper() else cleaned.title()
            if normalized.lower() in seen:
                continue
            seen.add(normalized.lower())
            found.append(normalized)
    return found


def _categorize_skills(skills: Sequence[str]) -> Dict[str, List[str]]:
    """Split skills into technical, soft, and other buckets using heuristics."""
    technical: List[str] = []
    soft: List[str] = []
    other: List[str] = []

    for raw_skill in _dedupe_preserve_order(skills):
        lowered = raw_skill.lower()
        if any(keyword in lowered for keyword in TECH_SKILL_HINTS) or re.search(r"[0-9+/]|\bapi\b", lowered):
            technical.append(raw_skill)
            continue
        if any(keyword in lowered for keyword in SOFT_SKILL_KEYWORDS):
            soft.append(raw_skill)
            continue
        # Heuristic: short uppercase abbreviations (e.g., PMP, PRINCE2) are likely certifications/technical
        if len(raw_skill) <= 5 and raw_skill.isupper():
            technical.append(raw_skill)
        else:
            other.append(raw_skill)

    return {
        "technical": technical,
        "soft": soft,
        "other": other,
    }


def _format_skills_section(skills: Dict[str, List[str]]) -> str:
    """Format categorized skills into a readable string."""
    parts: List[str] = []
    if skills.get("technical"):
        parts.append("Technical: " + ", ".join(skills["technical"]))
    if skills.get("soft"):
        parts.append("Soft: " + ", ".join(skills["soft"]))
    if skills.get("other"):
        parts.append("Additional: " + ", ".join(skills["other"]))
    return "\n".join(parts)


def _format_contact_section(contact: Dict[str, str]) -> Tuple[str, Dict[str, str]]:
    """Return formatted contact block string and sanitized contact dict."""
    sanitized = {
        "name": (contact.get("name") or "").strip(),
        "email": (contact.get("email") or "").strip(),
        "phone": (contact.get("phone") or "").strip(),
        "location": (contact.get("location") or "").strip(),
        "address": (contact.get("address") or contact.get("location") or "").strip(),
        "dob": (contact.get("date_of_birth") or contact.get("dob") or "").strip(),
        "linkedin": (contact.get("linkedin") or "").strip(),
        "github": (contact.get("github") or "").strip(),
        "website": (contact.get("website") or "").strip(),
    }
    if not sanitized["address"] and sanitized["location"]:
        sanitized["address"] = sanitized["location"]
    sanitized["date_of_birth"] = sanitized["dob"]

    lines = []
    if sanitized['name']:
        lines.append(f"Name: {sanitized['name']}")
    if sanitized['email']:
        lines.append(f"Email: {sanitized['email']}")
    if sanitized["phone"]:
        lines.append(f"Phone: {sanitized['phone']}")
    if sanitized.get("dob"):
        lines.append(f"DOB: {sanitized['dob']}")
    if sanitized.get("address"):
        lines.append(f"Address: {sanitized['address']}")
    elif sanitized["location"]:
        lines.append(f"Location: {sanitized['location']}")
    for field in ("linkedin", "github", "website"):
        value = sanitized.get(field)
        if value:
            lines.append(f"{field.title()}: {value}")
    return "\n".join(lines), sanitized


def _format_experience_section(experience: Sequence[Dict[str, str]]) -> str:
    """Format structured experience entries into ATS-friendly bullet lists."""
    entries: List[str] = []
    for job in experience:
        if not isinstance(job, dict):
            continue
        title = (job.get("title") or "Role").strip() or "Role"
        company = (job.get("company") or "Company").strip() or "Company"
        dates = (job.get("dates") or "").strip()
        location = (job.get("location") or "").strip()
        header_parts = [title, "|", company]
        header = " ".join(part for part in header_parts if part)
        if location:
            header = f"{header} ({location})"
        if dates:
            header = f"{header} — {dates}"

        raw_points = [p.strip() for p in job.get("points", []) if isinstance(p, str) and p.strip()]
        if raw_points:
            prepped = "\n".join(f"- {point}" for point in raw_points)
            strengthened = strengthen_experience_points(prepped)
            bullet_lines = [line for line in strengthened.splitlines() if line.strip()]
        else:
            bullet_lines = []

        section_lines = [header]
        section_lines.extend(bullet_lines)
        entries.append("\n".join(section_lines))

    return "\n\n".join(entries)


def _format_projects_section(projects: Sequence[Dict[str, str]]) -> str:
    blocks: List[str] = []
    for project in projects:
        if not isinstance(project, dict):
            continue
        name = (project.get("name") or "Project").strip()
        if not name:
            continue
        desc = (project.get("desc") or project.get("description") or "").strip()
        technologies = project.get("technologies") or project.get("tech") or []
        lines = [name]
        if desc:
            lines.append(f"- {desc}")
        if technologies:
            tech = ", ".join(_dedupe_preserve_order(str(t) for t in technologies))
            if tech:
                lines.append(f"- Tech: {tech}")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def _format_education_section(education: Sequence[Dict[str, str]]) -> str:
    entries: List[str] = []
    for edu in education:
        if not isinstance(edu, dict):
            continue
        degree = (edu.get("degree") or "Qualification").strip()
        school = (edu.get("school") or edu.get("institution") or "Institution").strip()
        year = (edu.get("year") or edu.get("date") or "").strip()
        line = degree
        if school:
            line = f"{line} — {school}"
        if year:
            line = f"{line} ({year})"
        entries.append(line)
    return "\n".join(entries)


def _format_list_section(items: Sequence[str], prefix: str = "- ") -> str:
    cleaned = _dedupe_preserve_order(str(item) for item in items)
    if not cleaned:
        return ""
    return "\n".join(f"{prefix}{item}" for item in cleaned)


def _extract_languages(sections: Dict[str, str]) -> List[str]:
    """Identify languages from any section content using simple heuristics."""
    candidates: List[str] = []
    search_space = "\n".join(str(v) for v in sections.values() if v)
    for match in re.finditer(r"languages?\s*[:\-]\s*([^\n]+)", search_space, flags=re.IGNORECASE):
        fragment = match.group(1)
        parts = re.split(r"[,;/]", fragment)
        candidates.extend(p.strip() for p in parts if p.strip())
    for name in LANGUAGE_NAMES:
        pattern = rf"\b{name}\b"
        if re.search(pattern, search_space, flags=re.IGNORECASE):
            candidates.append(name.title())
    return _dedupe_preserve_order(candidates)

def extract_text_from_pdf(file_stream):
    """Extract text from PDF using PyMuPDF and clean output."""
    try:
        # Use PyMuPDF for extraction
        if FITZ_AVAILABLE:
            pdf_bytes = file_stream.read()
            file_stream.seek(0)  # Reset stream
            
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                text_parts = []
                for page in doc:
                    page_text = page.get_text("text")
                    if page_text.strip():
                        text_parts.append(page_text)
                raw = "\n\n".join(text_parts)
        else:
            # Fallback to pdfminer
            raw = pdf_extract_text(file_stream)
        
        if not raw:
            return ""
        
        # Apply normalization and cleaning
        norm = normalize_text(raw)
        cleaned = _cleaner.clean_full_text(norm)
        return cleaned
    except Exception as e:
        logger.exception("PDF extraction failed: %s", e)
        return ""


def extract_text_from_any(file_bytes: bytes, filename: Optional[str]) -> str:
    """Extract text from supported CV formats (PDF, DOC, DOCX, raw text)."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(io.BytesIO(file_bytes))
    if name.endswith(".docx"):
        return extract_text_from_docx_bytes(file_bytes)
    if name.endswith(".doc"):
        return extract_text_from_doc_bytes(file_bytes)

    decoded = _safe_decode_bytes(file_bytes)
    return normalize_text(decoded)


def generate_pdf(text):
    """Generate professional PDF from CV text or structured data using WeasyPrint.
    
    Args:
        text: Either a string (plain CV text) or dict (structured CV data)
        
    Returns:
        BytesIO object containing the PDF
    """
    try:
        from xhtml2pdf import pisa
        from app.utils.cv_templates import build_professional_cv_html
        
        # Build HTML from input
        html_content = build_professional_cv_html(text)
        
        # Generate PDF with xhtml2pdf (Windows-compatible)
        pdf_io = io.BytesIO()
        pisa_status = pisa.CreatePDF(html_content, dest=pdf_io)
        
        if pisa_status.err:
            logger.warning("xhtml2pdf reported errors, falling back to FPDF")
            return _generate_pdf_fallback(text)
        
        # Return as BytesIO
        pdf_io.seek(0)
        return pdf_io
        
    except ImportError:
        logger.warning("xhtml2pdf not available, falling back to basic PDF generation")
        # Fallback to FPDF if xhtml2pdf not installed
        return _generate_pdf_fallback(text)
    except Exception as e:
        logger.error(f"PDF generation failed: {e}", exc_info=True)
        # Try fallback on any error
        return _generate_pdf_fallback(text)


def _generate_pdf_fallback(text):
    """Fallback PDF generation using FPDF (legacy)."""
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.add_page()
    font_path = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")
    if os.path.exists(font_path):
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)
    else:
        pdf.set_font("Helvetica", size=12)
        if isinstance(text, str):
            text = text.encode("latin-1", "replace").decode("latin-1")
    pdf.set_auto_page_break(auto=True, margin=15)

    def get_width(indent=0):
        return pdf.w - pdf.l_margin - pdf.r_margin - indent

    # If caller passed a dict, treat it as structured sections and render template
    if isinstance(text, dict):
        sections = text
        # Header placeholder
        header = sections.get("header") or "NAME\nEmail: you@example.com | Phone: +1-555-555-5555\nLocation: City, Country"
        for hline in header.split("\n"):
            pdf.set_font(pdf.font_family, "B", 14)
            pdf.cell(0, 8, hline, ln=True)
        pdf.ln(4)

        def render_section(title, content):
            if not content:
                return
            pdf.set_font(pdf.font_family, "B", 12)
            pdf.cell(0, 8, title, ln=True)
            pdf.set_font(pdf.font_family, size=11)
            # render bullets specially
            for line in str(content).split("\n"):
                line = line.strip()
                if not line:
                    pdf.ln(2)
                    continue
                if line.startswith("-"):
                    bullet = "•" if pdf.font_family == "DejaVu" else "-"
                    pdf.cell(6)
                    pdf.multi_cell(get_width(6), 6, f"{bullet} {line.lstrip('- ').strip()}")
                else:
                    pdf.multi_cell(get_width(), 6, line)
            pdf.ln(2)

        # Render in sensible order
        render_section("PROFESSIONAL SUMMARY", sections.get("about") or sections.get("summary"))
        render_section("KEY SKILLS", sections.get("skills"))
        render_section("WORK EXPERIENCE", sections.get("work_experience") or sections.get("experience"))
        render_section("PROJECTS", sections.get("projects"))
        render_section("ACHIEVEMENTS & EXTRACURRICULAR", sections.get("achievements") or sections.get("other"))
        render_section("EDUCATION", sections.get("education"))
    else:
        heading_prefixes = (
            "education",
            "experience",
            "skills",
            "projects",
            "contact",
            "professional summary",
            "key skills",
            "work experience",
            "certifications",
            "achievements",
            "languages",
            "additional information",
        )
        for line in str(text).split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue
            if line.lower().startswith(heading_prefixes):
                pdf.set_font(pdf.font_family, "B", 14)
                pdf.cell(0, 10, line, ln=True)
                pdf.set_font(pdf.font_family, size=12)
            elif line.startswith(("-", "*")) or (line and line[0:2].isdigit()):
                indent = 10
                pdf.cell(indent)
                bullet = "•" if pdf.font_family == "DejaVu" else "-"
                safe_text = line.lstrip("-*0123456789. ")
                pdf.multi_cell(get_width(indent), 8, f"{bullet} {safe_text}")
            else:
                pdf.multi_cell(get_width(), 8, line)
    # Use 'S' to get PDF as bytes (FPDF2 returns bytes directly)
    pdf_output = pdf.output(dest='S')
    
    # Handle both old FPDF (returns str) and FPDF2 (returns bytes)
    if isinstance(pdf_output, str):
        pdf_data = pdf_output.encode('latin-1')
    else:
        pdf_data = pdf_output  # Already bytes
    
    pdf_bytes = io.BytesIO(pdf_data)
    pdf_bytes.seek(0)
    return pdf_bytes


def _score_by_keywords(text, domain):
    """Return (score_fraction, missing_keywords, found_keywords) based on domain keywords."""
    if not domain:
        return 0.0, [], []
    domain = domain.lower().replace(" ", "_")
    keywords = DOMAIN_KEYWORDS.get(domain, [])
    text_lower = text.lower()
    found = [k for k in keywords if k.lower() in text_lower]
    missing = [k for k in keywords if k.lower() not in text_lower]
    score = 0.0
    if keywords:
        score = min(1.0, len(found) / len(keywords))
    return score, missing, found


def compute_ats_score(text, domain=None):
    """Compute a comprehensive ATS score (0-100).

    Scoring breakdown:
    - Contact Information (15 points): email, phone, location
    - Professional Summary (10 points): presence and quality
    - Skills Section (10 points): clear skills section
    - Work Experience (15 points): structured with dates and bullets
    - Education (10 points): degree and institution
    - Keywords (20 points): domain-specific keywords
    - Action Verbs (8 points): strong action verbs in experience
    - Quantifiable Achievements (7 points): numbers/metrics
    - Formatting (5 points): proper structure and sections
    - Length (5 points): appropriate length (300-1200 words)
    """
    score = 0
    text_lower = text.lower()
    
    # 1. Contact Information (15 points)
    contact_score = 0
    if re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text_lower):
        contact_score += 7
    if re.search(r"\+?\d[\d \-()]{7,}\d", text_lower):
        contact_score += 5
    if re.search(r"\b(city|state|country|location)\b.*?[,\n]", text_lower):
        contact_score += 3
    score += contact_score
    
    # 2. Professional Summary (10 points)
    summary_keywords = r"\b(summary|objective|profile|about)\b"
    if re.search(summary_keywords, text_lower):
        summary_match = re.search(rf"{summary_keywords}[:\s]*([^\n]*(?:\n[^\n]*)?)", text_lower, re.IGNORECASE)
        if summary_match and len(summary_match.group(1).split()) > 15:
            score += 10
        else:
            score += 5
    
    # 3. Skills Section (10 points)
    if re.search(r"\b(skills|competencies|expertise)\b", text_lower):
        skills_section = re.search(r"\b(skills|competencies)\b[^\n]*\n([^\n]*(?:\n[^\n]*){1,10})", text_lower, re.IGNORECASE)
        if skills_section:
            skills_text = skills_section.group(2)
            skill_count = len(re.findall(r"[,•\-]|\n", skills_text))
            if skill_count >= 5:
                score += 10
            else:
                score += 5
    
    # 4. Work Experience (15 points)
    experience_keywords = r"\b(experience|employment|work history)\b"
    if re.search(experience_keywords, text_lower):
        exp_score = 5
        # Check for dates
        if re.search(r"(20|19)\d{2}\s*[-–]\s*(20|19)?\d{0,4}|present|current", text_lower):
            exp_score += 5
        # Check for bullet points
        bullet_count = len(re.findall(r"^\s*[-•*]\s+", text, re.MULTILINE))
        if bullet_count >= 3:
            exp_score += 5
        score += exp_score
    
    # 5. Education (10 points)
    education_keywords = r"\b(education|academic|degree|university|college)\b"
    if re.search(education_keywords, text_lower):
        edu_score = 5
        # Check for degree keywords
        if re.search(r"\b(bachelor|master|phd|b\.?s\.?|m\.?s\.?|b\.?a\.?|m\.?a\.|diploma)\b", text_lower):
            edu_score += 5
        score += edu_score
    
    # 6. Domain Keywords (20 points)
    kw_score, missing, found = _score_by_keywords(text, domain)
    score += int(20 * kw_score)
    
    # 7. Action Verbs (8 points)
    action_verb_count = 0
    for verb in ACTION_VERBS:
        if verb in text_lower:
            action_verb_count += 1
    action_score = min(8, action_verb_count)
    score += action_score
    
    # 8. Quantifiable Achievements (7 points)
    # Look for numbers, percentages, metrics
    metrics_pattern = r"\d+[%+]|\$\d+|\d+\s*(users|clients|customers|projects|team|members|revenue|sales|growth)"
    metrics_count = len(re.findall(metrics_pattern, text_lower))
    score += min(7, metrics_count * 2)
    
    # 9. Formatting (5 points)
    # Check for proper section structure
    section_count = sum(1 for kw in ["summary", "experience", "education", "skills"] if kw in text_lower)
    score += min(5, section_count)
    
    # 10. Length (5 points)
    words = len(re.findall(r"\w+", text))
    if 300 <= words <= 1200:
        score += 5
    elif 200 <= words < 300 or 1200 < words <= 1500:
        score += 3
    elif words >= 150:
        score += 1
    
    return min(100, score), missing, found


def analyze_ats_score_detailed(text, domain=None):
    """Optimized ATS analysis with comprehensive scoring.
    
    Returns dict with:
    - overall_score: 0-100
    - breakdown: dict of category scores
    - missing_elements: list of what's missing
    - recommendations: list of improvement suggestions
    - missing_keywords: keywords to add for domain
    - found_keywords: keywords already present
    - strengths: list of CV strengths
    - category_scores: normalized scores by category
    """
    text_lower = text.lower()
    breakdown = {}
    missing_elements = []
    recommendations = []
    
    # Precompile patterns for performance
    email_pattern = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
    phone_pattern = re.compile(r"\+?\d[\d \-()]{7,}\d")
    location_pattern = re.compile(r"\b(city|state|country|location|address)\b", re.I)
    date_pattern = re.compile(r"(20|19)\d{2}\s*[-–]\s*(20|19)?\d{0,4}|present|current", re.I)
    
    # Contact Information (15 points)
    contact_score = 0
    has_email = email_pattern.search(text_lower)
    has_phone = phone_pattern.search(text_lower)
    has_location = location_pattern.search(text_lower)
    
    if has_email: 
        contact_score += 7
    else: 
        missing_elements.append("Email address")
        recommendations.append("⚠️ Add your professional email address")
    
    if has_phone: 
        contact_score += 5
    else: 
        missing_elements.append("Phone number")
        recommendations.append("📞 Include contact phone number")
    
    if has_location: 
        contact_score += 3
    else: 
        recommendations.append("📍 Add your location (city/country)")
    
    breakdown["Contact Information"] = contact_score
    
    # Professional Summary (12 points)
    summary_score = 0
    summary_pattern = re.compile(r"\b(summary|objective|profile|about|professional\s+summary)\b[:\s]*(.{50,500})", re.I | re.DOTALL)
    summary_match = summary_pattern.search(text)
    
    if summary_match:
        summary_text = summary_match.group(2)
        word_count = len(summary_text.split())
        if word_count >= 30:
            summary_score = 12
        elif word_count >= 15:
            summary_score = 8
            recommendations.append("💡 Expand professional summary to 30-50 words")
        else:
            summary_score = 4
            recommendations.append("📝 Professional summary is too brief - aim for 30-50 words")
    else:
        missing_elements.append("Professional Summary")
        recommendations.append("🎯 Add a professional summary highlighting your key strengths and career goals")
    
    breakdown["Professional Summary"] = summary_score
    
    # Skills (15 points)
    skills_score = 0
    skills_pattern = re.compile(r"\b(skills|competencies|expertise|technical\s+skills|core\s+competencies)\b", re.I)
    has_skills = skills_pattern.search(text_lower)
    
    if has_skills:
        # Count skills by looking for delimiters and keywords
        skill_delimiters = len(re.findall(r"[,•\-\|]", text))
        technical_terms = len(re.findall(r"\b(python|java|javascript|sql|aws|docker|react|node|api|database|cloud|agile|scrum)\b", text_lower))
        
        total_skill_indicators = skill_delimiters + technical_terms
        
        if total_skill_indicators >= 12:
            skills_score = 15
        elif total_skill_indicators >= 8:
            skills_score = 12
        elif total_skill_indicators >= 5:
            skills_score = 8
            recommendations.append("💼 Add more technical and soft skills (aim for 10-15)")
        else:
            skills_score = 4
            recommendations.append("📊 Skills section needs expansion with relevant keywords")
    else:
        missing_elements.append("Skills Section")
        recommendations.append("⚡ Create a dedicated Skills section with 10-15 relevant skills")
    
    breakdown["Skills"] = skills_score
    
    # Work Experience (20 points)
    exp_score = 0
    exp_pattern = re.compile(r"\b(experience|employment|work\s+history|professional\s+experience|career)\b", re.I)
    has_experience = exp_pattern.search(text_lower)
    
    if has_experience:
        exp_score = 6
        
        # Check for dates
        dates_found = date_pattern.findall(text)
        if len(dates_found) >= 2:
            exp_score += 5
        elif len(dates_found) >= 1:
            exp_score += 3
            recommendations.append("📅 Add dates to all work experience entries")
        else:
            recommendations.append("⏰ Include employment dates for all positions")
        
        # Check for bullet points (achievements)
        bullet_count = len(re.findall(r"^\s*[-•*]\s+", text, re.MULTILINE))
        if bullet_count >= 6:
            exp_score += 6
        elif bullet_count >= 3:
            exp_score += 4
            recommendations.append("✨ Add more bullet points describing achievements")
        else:
            exp_score += 1
            recommendations.append("🔸 Use bullet points to highlight key accomplishments")
        
        # Check for job titles and company names
        job_indicators = len(re.findall(r"\b(manager|developer|engineer|analyst|coordinator|specialist|director|lead|senior|junior)\b", text_lower))
        if job_indicators >= 2:
            exp_score += 3
    else:
        missing_elements.append("Work Experience")
        recommendations.append("💼 Add Work Experience section with job titles, companies, and achievements")
    
    breakdown["Work Experience"] = exp_score
    
    # Education (12 points)
    edu_score = 0
    edu_pattern = re.compile(r"\b(education|academic|qualification|degree|university|college|institute)\b", re.I)
    has_education = edu_pattern.search(text_lower)
    
    if has_education:
        edu_score = 6
        
        # Check for degree type
        degree_pattern = re.compile(r"\b(bachelor|master|phd|doctorate|b\.?s\.?c?|m\.?s\.?c?|b\.?a\.|m\.?a\.|diploma|associate)\b", re.I)
        if degree_pattern.search(text_lower):
            edu_score += 4
        else:
            recommendations.append("🎓 Specify your degree type (e.g., Bachelor of Science in Computer Science)")
        
        # Check for institution names
        if re.search(r"\b(university|college|institute|school)\b", text_lower):
            edu_score += 2
    else:
        missing_elements.append("Education")
        recommendations.append("🎓 Add Education section with degree, institution, and graduation year")
    
    breakdown["Education"] = edu_score
    
    # Domain Keywords (15 points)
    kw_score, missing_kw, found_kw = _score_by_keywords(text, domain)
    breakdown["Domain Keywords"] = int(15 * kw_score)
    if len(missing_kw) > 0:
        recommendations.append(f"🔑 Add relevant keywords: {', '.join(missing_kw[:5])}")
    elif kw_score < 0.5:
        recommendations.append("🔍 Include more industry-specific keywords and technologies")
    
    # Action Verbs (10 points)
    action_count = sum(1 for verb in ACTION_VERBS if f" {verb} " in f" {text_lower} " or f" {verb}ed " in f" {text_lower} ")
    action_score = min(10, int(action_count * 1.5))
    breakdown["Action Verbs"] = action_score
    
    if action_count < 3:
        recommendations.append("💪 Use more strong action verbs (achieved, led, developed, managed, implemented)")
    elif action_count < 6:
        recommendations.append("💡 Add more action verbs to strengthen impact statements")
    
    # Quantifiable Achievements (10 points)
    metrics_pattern = re.compile(r"\d+[%+]|\$\d+k?m?|\d+\s*(users|clients|customers|projects|team|members|employees|revenue|sales|growth|increase|decrease|reduction|improvement)", re.I)
    metrics_matches = metrics_pattern.findall(text_lower)
    metrics_count = len(metrics_matches)
    
    achievement_score = min(10, metrics_count * 2)
    breakdown["Quantifiable Achievements"] = achievement_score
    
    if metrics_count == 0:
        recommendations.append("📊 Add quantifiable achievements with numbers/percentages (e.g., 'Increased sales by 25%')")
    elif metrics_count < 3:
        recommendations.append("📈 Include more metrics to demonstrate measurable impact")
    
    # Document Structure (6 points)
    section_keywords = ["summary", "objective", "experience", "employment", "education", "skills", "projects", "certifications"]
    section_count = sum(1 for kw in section_keywords if kw in text_lower)
    
    structure_score = min(6, int(section_count * 1.2))
    breakdown["Document Structure"] = structure_score
    
    if section_count < 4:
        recommendations.append("📋 Add more standard sections (Summary, Skills, Experience, Education, Projects)")
    
    # Formatting & Length (7 points)
    words = len(re.findall(r"\w+", text))
    length_score = 0
    
    if 400 <= words <= 800:
        length_score = 7
    elif 300 <= words < 400 or 800 < words <= 1000:
        length_score = 5
        if words < 400:
            recommendations.append("📝 Add more detail - optimal length is 400-800 words")
    elif 200 <= words < 300:
        length_score = 3
        recommendations.append("📄 CV is too brief - expand with more details (aim for 400-800 words)")
    elif words > 1000:
        length_score = 4
        recommendations.append("✂️ Consider condensing to 1-2 pages for better ATS readability")
    else:
        length_score = 1
        recommendations.append("⚠️ CV is too short - add substantial content")
    
    breakdown["Formatting & Length"] = length_score
    
    overall_score = sum(breakdown.values())
    
    # Generate strengths list based on scores
    strengths = []
    if breakdown.get("Contact Information", 0) >= 13:
        strengths.append("✅ Complete and professional contact details")
    if breakdown.get("Professional Summary", 0) >= 10:
        strengths.append("✅ Strong and compelling professional summary")
    if breakdown.get("Skills", 0) >= 12:
        strengths.append("✅ Comprehensive skills section with relevant keywords")
    if breakdown.get("Work Experience", 0) >= 16:
        strengths.append("✅ Well-documented work experience with clear achievements")
    if breakdown.get("Education", 0) >= 10:
        strengths.append("✅ Complete educational background")
    if breakdown.get("Quantifiable Achievements", 0) >= 6:
        strengths.append("✅ Strong use of metrics and quantifiable results")
    if breakdown.get("Action Verbs", 0) >= 7:
        strengths.append("✅ Effective use of action verbs and impact statements")
    if breakdown.get("Domain Keywords", 0) >= 10:
        strengths.append("✅ Good keyword optimization for ATS systems")
        
    # Priority recommendations based on score tiers
    priority_recs = []
    if overall_score < 50:
        priority_recs.append("🎯 URGENT: CV needs major improvements to pass ATS screening")
        priority_recs.append("🎯 Focus on: Complete sections, add keywords, include achievements with metrics")
    elif overall_score < 70:
        priority_recs.append("⚡ IMPORTANT: Enhance CV to improve ATS ranking")
        priority_recs.append("💡 Focus on: More keywords, quantifiable achievements, and detailed experience")
    elif overall_score < 85:
        priority_recs.append("✨ Good foundation - refine details for maximum impact")
    
    # Calculate category scores (for frontend visualization)
    category_scores = {
        "contact": int((breakdown.get("Contact Information", 0) / 15) * 100),
        "content": int(((breakdown.get("Professional Summary", 0) + breakdown.get("Work Experience", 0)) / 32) * 100),
        "keywords": int(((breakdown.get("Skills", 0) + breakdown.get("Domain Keywords", 0)) / 30) * 100),
        "impact": int(((breakdown.get("Action Verbs", 0) + breakdown.get("Quantifiable Achievements", 0)) / 20) * 100),
        "structure": int(((breakdown.get("Document Structure", 0) + breakdown.get("Formatting & Length", 0)) / 13) * 100),
    }
    
    return {
        "overall_score": min(100, overall_score),
        "breakdown": breakdown,
        "category_scores": category_scores,
        "missing_elements": missing_elements,
        "recommendations": priority_recs + recommendations,
        "strengths": strengths,
        "missing_keywords": missing_kw if domain else [],
        "found_keywords": found_kw if domain else [],
        "grade": "Excellent" if overall_score >= 85 else "Good" if overall_score >= 70 else "Fair" if overall_score >= 50 else "Needs Improvement"
    }


def extract_sections(text):
    """Split text into structured sections using heading-aware heuristics and NLP."""
    sections: Dict[str, str] = {key: "" for key in SECTION_KEYS}

    if not text:
        return sections

    lines = text.split('\n')
    current_key = "about"
    buffer: List[str] = []

    def _commit_buffer():
        if not buffer:
            return
        block = "\n".join(buffer).strip()
        if not block:
            buffer.clear()
            return
        existing = sections.get(current_key, "")
        sections[current_key] = (existing + "\n" + block).strip() if existing else block
        buffer.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if buffer and buffer[-1] != "":
                buffer.append("")
            continue

        candidate = line
        candidate_is_bullet = False
        if candidate[0] in BULLET_PREFIXES:
            candidate = candidate.lstrip(''.join(BULLET_PREFIXES) + ' ')
            candidate = candidate.strip()
            candidate_is_bullet = True

        heading_key: Optional[str] = None
        inline_body: Optional[str] = None
        
        # 1. Try Regex/Exact Match
        inline_match = INLINE_HEADING_PATTERN.match(candidate)
        if inline_match:
            heading_key = _heading_lookup(inline_match.group("label"))
            inline_body = inline_match.group("body").strip()
        if not heading_key:
            heading_key = _heading_lookup(candidate)
            
        # 2. Try NLP Classification if regex failed and line is short enough to be a header
        if not heading_key and len(candidate.split()) <= 6 and not candidate_is_bullet:
             heading_key = classify_header_nlp(candidate, SECTION_SYNONYMS)

        if heading_key and (not candidate_is_bullet or len(candidate.split()) <= 5):
            _commit_buffer()
            current_key = heading_key
            if inline_body:
                buffer.append(inline_body)
            continue

        buffer.append(line)

    _commit_buffer()

    _augment_sections_from_keywords(text, sections)

    for key in sections:
        sections[key] = re.sub(r'\n{2,}', '\n\n', sections[key].strip())

    sections["summary"] = sections.get("about", sections.get("summary", ""))

    return sections


def build_standardized_sections(cv_text: str) -> Dict[str, object]:
    """Return structured sections aligned with the standardized preview spec."""
    normalized_text = normalize_text(cv_text or "")
    raw_sections = extract_sections(normalized_text)

    about_text = raw_sections.get("about", "")
    contact_info = extract_contact_info(about_text)
    contact_block, sanitized_contact = _format_contact_section(contact_info)

    # Remove lines that contain contact details from summary
    summary_lines: List[str] = []
    for line in about_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", stripped):
            continue
        if re.search(r"\+?[\d\s\-()]{7,}", stripped):
            continue
        lowered = stripped.lower()
        if any(keyword in lowered for keyword in ("email", "phone", "mobile", "linkedin", "github")):
            continue
        if DOB_LABEL_PATTERN.search(lowered):
            continue
        if sanitized_contact.get("dob") and sanitized_contact["dob"] in stripped:
            continue
        if sanitized_contact.get("address") and sanitized_contact["address"] in stripped:
            continue
        summary_lines.append(stripped)
    summary_text = " ".join(summary_lines).strip()

    skills_candidates: List[str] = []
    raw_skills = raw_sections.get("skills", "")
    if raw_skills:
        skills_candidates = [s.strip() for s in re.split(r"[,;\n]", raw_skills) if s.strip()]
    if not skills_candidates:
        skills_candidates = _infer_skills_from_text(normalized_text)
    categorized_skills = _categorize_skills(skills_candidates)
    skills_formatted = _format_skills_section(categorized_skills)

    experience_structured = parse_experience_section(raw_sections.get("experience", ""))
    experience_formatted = _format_experience_section(experience_structured)

    projects_structured = parse_projects_section(raw_sections.get("projects", ""))
    projects_formatted = _format_projects_section(projects_structured)

    education_structured = parse_education_section(raw_sections.get("education", ""))
    education_formatted = _format_education_section(education_structured)

    volunteer_structured = parse_experience_section(raw_sections.get("volunteer", ""))

    certifications_list: List[str] = []
    achievements_list: List[str] = []

    for line in _split_section_lines(raw_sections.get("certifications", "")):
        certifications_list.append(line)

    for source in (raw_sections.get("achievements", ""), raw_sections.get("other", "")):
        if not source:
            continue
        for line in _split_section_lines(source):
            lowered = line.lower()
            if any(keyword in lowered for keyword in ("cert", "license", "licence", "credential", "honor", "award")):
                certifications_list.append(line)
            else:
                achievements_list.append(line)

    languages: List[str] = []
    if raw_sections.get("languages"):
        for chunk in re.split(r"[,/;\n]", raw_sections["languages"]):
            cleaned = chunk.strip()
            if not cleaned:
                continue
            languages.append(cleaned.title() if cleaned.islower() else cleaned)
    languages.extend(_extract_languages(raw_sections))
    languages = _dedupe_preserve_order(languages)

    additional_text = raw_sections.get("other", "") or ""
    additional_text = additional_text.strip()

    structured = {
        "contact_information": {**sanitized_contact, "block": contact_block},
        "professional_summary": summary_text,
        "skills": {**categorized_skills, "formatted": skills_formatted},
        "work_experience": experience_structured,
        "projects": projects_structured,
        "education": education_structured,
        "certifications": _dedupe_preserve_order(certifications_list),
        "achievements": _dedupe_preserve_order(achievements_list),
        "languages": languages,
        "volunteer_experience": volunteer_structured,
        "additional_information": additional_text,
    }

    # Augment findings with NLP if available (noun chunks -> skills, entities -> names/orgs)
    try:
        from app.utils import nlp_utils
        nlp = nlp_utils.load_spacy_model()
        if nlp:
            try:
                ents = nlp_utils.extract_entities(cv_text)
                noun_chunks = nlp_utils.extract_noun_chunks(cv_text)
                # Add named entities to contact info when missing
                if not structured['contact_information'].get('name') and ents.get('PERSON'):
                    structured['contact_information']['name'] = ents['PERSON'][0]
                if not structured['contact_information'].get('location') and ents.get('GPE'):
                    structured['contact_information']['location'] = ents['GPE'][0]
                # Augment skill candidates with noun chunks heuristically
                inferred_skills = [nc for nc in noun_chunks if 2 <= len(nc.split()) <= 4]
                # merge into skills 'technical' if not already present
                existing_all = []
                if isinstance(structured.get('skills'), dict):
                    existing_all = structured['skills'].get('all') or []
                merged = _dedupe_preserve_order(list(existing_all) + inferred_skills)
                if isinstance(structured.get('skills'), dict):
                    structured['skills']['all'] = merged
            except Exception:
                logger.debug("NLP augmentation failed, continuing without it")
    except Exception:
        # nlp_utils may not be present or spaCy not installed - that's fine
        pass

    return structured


def _structured_to_preview(structured: Dict[str, object]) -> Tuple[List[Dict[str, str]], Dict[str, str], str]:
    """Convert structured sections into ordered preview sections and optimized text."""
    ordered_sections: List[Dict[str, str]] = []
    sections_map: Dict[str, str] = {}

    # Helper to format contact info for the top block
    contact = structured.get("contact_information", {}) if isinstance(structured.get("contact_information"), dict) else {}
    name = contact.get("name") or ""
    # We don't add contact to ordered_sections loop in the same way for the text output
    # but we keep it in ordered_sections for the frontend UI if it needs it.
    
    # We will build the optimized_text separately to strictly follow the user's format
    text_lines = []
    
    # 1. Header
    if name:
        text_lines.append(f"**{name}**")
    # Job title is hard to guess from rule-based, skip or try to find latest role
    
    if contact.get("phone"): text_lines.append(f"Phone: {contact['phone']}")
    if contact.get("email"): text_lines.append(f"Email: {contact['email']}")
    if contact.get("linkedin"): text_lines.append(f"LinkedIn: {contact['linkedin']}")
    if contact.get("github"): text_lines.append(f"GitHub: {contact['github']}")
    if contact.get("website"): text_lines.append(f"Portfolio: {contact['website']}")
    
    text_lines.append("")
    text_lines.append("---")
    text_lines.append("")

    for key, label in STANDARD_SECTION_ORDER:
        content = ""
        if key == "contact_information":
            # Already handled in header for text, but we add to ordered_sections for UI
            block = structured.get(key, {}).get("block") if isinstance(structured.get(key), dict) else ""
            content = block
        elif key == "professional_summary":
            summary = structured.get(key) or ""
            content = summary
        elif key == "skills":
            skills_block = ""
            if isinstance(structured.get(key), dict):
                skills_block = structured[key].get("formatted") or ""
            content = skills_block
        elif key == "work_experience":
            content = _format_experience_section(structured.get(key, [])) if structured.get(key) else ""
        elif key == "projects":
            content = _format_projects_section(structured.get(key, [])) if structured.get(key) else ""
        elif key == "education":
            content = _format_education_section(structured.get(key, [])) if structured.get(key) else ""
        elif key in ("certifications", "achievements", "languages"):
            items = structured.get(key) or []
            content = _format_list_section(items) if items else ""
        elif key == "volunteer_experience":
            content = _format_experience_section(structured.get(key, [])) if structured.get(key) else ""
        elif key == "additional_information":
            addl = structured.get(key) or ""
            content = addl

        clean_content = content.strip()
        if not clean_content:
            continue

        sections_map[key] = clean_content
        ordered_sections.append({
            "key": key,
            "label": label,
            "content": clean_content,
        })
        
        # Add to text output (skip contact info as it's already at top)
        if key != "contact_information":
            text_lines.append(f"## {label.upper()}")
            text_lines.append(clean_content)
            text_lines.append("")

    optimized_text = "\n".join(text_lines).strip()
    return ordered_sections, sections_map, optimized_text


def _structured_to_legacy_sections(structured: Dict[str, object]) -> Dict[str, str]:
    """Create a legacy sections dict compatible with existing template utilities."""
    contact = structured.get("contact_information", {}) if isinstance(structured.get("contact_information"), dict) else {}
    summary = structured.get("professional_summary") or ""
    summary_block = "\n".join(line for line in (contact.get("block"), summary) if line)

    skills_section = structured.get("skills", {}) if isinstance(structured.get("skills"), dict) else {}
    all_skills: List[str] = []
    for bucket in ("technical", "soft", "other"):
        all_skills.extend(skills_section.get(bucket) or [])
    skills_text = ", ".join(_dedupe_preserve_order(all_skills))

    experience_text = _format_experience_section(structured.get("work_experience", []))
    projects_text = _format_projects_section(structured.get("projects", []))
    education_text = _format_education_section(structured.get("education", []))

    achievements_combo = _dedupe_preserve_order(
        list(structured.get("certifications", []) or []) + list(structured.get("achievements", []) or [])
    )
    if structured.get("languages"):
        achievements_combo.extend([f"Language: {lang}" for lang in structured.get("languages", [])])
    for volunteer in structured.get("volunteer_experience", []) or []:
        if isinstance(volunteer, dict):
            title = (volunteer.get("title") or "Volunteer").strip()
            organization = (volunteer.get("company") or volunteer.get("organization") or "Organization").strip()
            dates = (volunteer.get("dates") or "").strip()
            entry = f"Volunteer: {title} at {organization}"
            if dates:
                entry = f"{entry} ({dates})"
            achievements_combo.append(entry)

    achievements_text = "\n".join(_dedupe_preserve_order(achievements_combo))

    return {
        "about": summary_block,
        "skills": skills_text,
        "experience": experience_text,
        "projects": projects_text,
        "education": education_text,
        "achievements": achievements_text,
        "other": structured.get("additional_information") or "",
    }


ACTION_VERBS = [
    "led", "built", "designed", "developed", "implemented", "improved", "optimized", "created",
    "reduced", "increased", "managed", "launched", "orchestrated", "analyzed", "automated", "mentored"
]


def normalize_bullets(text):
    """Ensure bullets are on separate lines, start with action verbs where possible."""
    lines = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        # turn commas in long sentences into bullets if looks like list
        if "," in line and len(line.split()) > 12 and not line.startswith("-"):
            parts = [p.strip() for p in line.split(",") if p.strip()]
            for p in parts:
                lines.append(f"- {p}")
            continue
        # ensure bullet prefix
        if line[0] in "-*•" or re.match(r"^\d+\.", line):
            # ensure starts with action verb
            content = re.sub(r"^[\-*•\s\d.]+", "", line).strip()
            first_word = content.split()[0].lower() if content else ""
            if first_word not in ACTION_VERBS and content:
                content = ACTION_VERBS[0] + " " + content
            lines.append(f"- {content}")
        else:
            lines.append(line)
    return "\n".join(lines)


def strengthen_experience_points(text):
    """Ensure experience bullets start with action verbs and are concise.

    This is a lightweight, rule-based improvement applied when AI is unavailable.
    """
    if not text:
        return text
    out_lines = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            out_lines.append("")
            continue
        # if it's a bullet, ensure it begins with an action verb
        is_bullet = line.startswith("-") or line.startswith("*") or line.startswith("•")
        content = line.lstrip('-*• ').strip()
        words = content.split()
        if words:
            first = words[0].lower()
            if first not in ACTION_VERBS:
                # Prepend a sensible action verb
                content = ACTION_VERBS[0].capitalize() + ' ' + content
            else:
                # capitalize first word
                content = words[0].capitalize() + ' ' + ' '.join(words[1:])

        if is_bullet:
            out_lines.append(f"- {content}")
        else:
            out_lines.append(content)

    return '\n'.join(out_lines)


def optimize_cv_rule_based(cv_text: str, job_domain: Optional[str] = None) -> Dict[str, object]:
    """Produce a cleaned, ATS-friendly CV structure without relying on AI."""
    structured = build_standardized_sections(cv_text or "")
    ordered_sections, sections_map, optimized_text = _structured_to_preview(structured)
    structured_payload = build_structured_cv_payload(structured)

    # Compute ATS score and keyword coverage
    ats_score, missing_keywords, found_keywords = compute_ats_score(optimized_text, job_domain)

    suggestions = _generate_suggestions(structured, missing_keywords)

    legacy_sections = _structured_to_legacy_sections(structured)
    template_data = convert_to_template_format(legacy_sections)
    extracted = build_extracted_sections(cv_text, structured_sections=structured)

    return {
        "optimized_text": optimized_text,
        "optimized_cv": optimized_text,
        "optimized_ats_cv": optimized_text,
        "sections": sections_map,
        "ordered_sections": ordered_sections,
        "structured_sections": structured,
        "structured_cv": structured_payload,
        "extracted": extracted,
        "extracted_text": cv_text,
        "suggestions": suggestions,
        "ats_score": ats_score,
        "recommended_keywords": missing_keywords,
        "found_keywords": found_keywords,
        "template_data": template_data,
    }

def parse_experience_section(text):
    """Parse experience section text into structured format.
    
    Handles multiple formats:
    - Job Title at Company (Date Range)
    - Company | Job Title | Dates
    - Date Range: Job Title at Company
    
    Returns list of {title, company, dates, points[]}.
    """
    if not text or not text.strip():
        return []
    
    jobs = []
    current_job = None
    lines = text.split('\n')
    
    # Regex patterns for date ranges
    date_pattern = r'\(([^)]*(?:20|19)\d{2}[^)]*)\)|\[([^\]]*(?:20|19)\d{2}[^\]]*)\]|(\w+\s+\d{4}\s*[-–]\s*(?:\w+\s+)?\d{4}|Present|Current)'
    
    for line in lines:
        line = line.strip()
        if not line:
            if current_job and current_job.get('points'):
                jobs.append(current_job)
                current_job = None
            continue
        
        # Check if this is a new job entry (non-bulleted, contains company/title info)
        if not line.startswith(('-', '•', '*')) and any(sep in line for sep in [' at ', ' | ', ' - ', '–']):
            # Save previous job
            if current_job and (current_job.get('points') or current_job.get('title')):
                jobs.append(current_job)
            
            # Parse new job
            title = ""
            company = ""
            dates = ""
            
            # Extract dates first
            date_match = re.search(date_pattern, line, re.IGNORECASE)
            if date_match:
                dates = date_match.group(1) or date_match.group(2) or date_match.group(3)
                line_without_dates = re.sub(date_pattern, '', line).strip()
            else:
                line_without_dates = line
            
            # Parse title and company
            if ' at ' in line_without_dates:
                title, company = [x.strip() for x in line_without_dates.split(' at ', 1)]
            elif ' | ' in line_without_dates:
                parts = [x.strip() for x in line_without_dates.split('|')]
                if len(parts) == 2:
                    company, title = parts
                elif len(parts) == 3:
                    company, title, dates_alt = parts
                    if not dates:
                        dates = dates_alt
            elif ' - ' in line_without_dates or '–' in line_without_dates:
                sep = ' – ' if '–' in line_without_dates else ' - '
                parts = [x.strip() for x in line_without_dates.split(sep, 1)]
                if len(parts) == 2 and any(year in parts[1] for year in ['20', '19', 'Present']):
                    title = parts[0]
                    dates = parts[1]
                else:
                    title = line_without_dates
            else:
                title = line_without_dates
            
            current_job = {
                'title': title or 'Position',
                'company': company or 'Company',
                'dates': dates or 'Present',
                'points': []
            }
        elif current_job and (line.startswith(('-', '•', '*')) or (current_job and current_job.get('points'))):
            # Bullet point or continuation
            point = line.lstrip('-•* ').strip()
            if point:
                current_job['points'].append(point)
        elif not current_job and line and any(sep in line for sep in [' at ', ' | ', ' - ', '–']):
            # Start new job entry
            title = ""
            company = ""
            dates = ""
            
            date_match = re.search(date_pattern, line, re.IGNORECASE)
            if date_match:
                dates = date_match.group(1) or date_match.group(2) or date_match.group(3)
                line_without_dates = re.sub(date_pattern, '', line).strip()
            else:
                line_without_dates = line
            
            if ' at ' in line_without_dates:
                title, company = [x.strip() for x in line_without_dates.split(' at ', 1)]
            elif ' | ' in line_without_dates:
                parts = [x.strip() for x in line_without_dates.split('|')]
                if len(parts) >= 2:
                    company = parts[0]
                    title = parts[1]
            
            current_job = {
                'title': title or 'Position',
                'company': company or 'Company',
                'dates': dates or 'Present',
                'points': [line] if line.startswith(('-', '•', '*')) else []
            }
    
    # Add last job if it has content
    if current_job and (current_job.get('points') or current_job.get('title')):
        jobs.append(current_job)
    
    # Filter out empty jobs and return
    return [j for j in jobs if j.get('title') and j.get('company')]


def parse_education_section(text):
    """Parse education section into structured format.
    
    Handles formats:
    - Degree - University (Year)
    - University | Degree | Year
    - Degree (Year)
    - Bachelor of Science in Computer Science - MIT (2020)
    
    Returns list of {degree, school, year}.
    """
    if not text or not text.strip():
        return []
    
    education = []
    year_pattern = r'\(([^)]*(?:20|19)\d{2}[^)]*)\)|\[([^\]]*)\]|(?:^|\s)(\d{4})(?:\s|$)|(?:graduating|graduated|expected)\s+(\d{4}|present)'
    
    for line in text.split('\n'):
        line = line.strip()
        if not line or line.startswith(('-', '•', '*')):
            # Skip bullets - they might be details
            if line and line.startswith(('-', '•', '*')):
                point = line.lstrip('-•* ').strip()
                # Only process if looks like an achievement detail
                continue
            continue
        
        degree = ""
        school = ""
        year = ""
        
        # Extract year/date
        year_match = re.search(year_pattern, line, re.IGNORECASE)
        if year_match:
            year = year_match.group(1) or year_match.group(2) or year_match.group(3) or year_match.group(4)
            line_without_year = re.sub(year_pattern, '', line, flags=re.IGNORECASE).strip()
        else:
            line_without_year = line
        
        # Split by common separators
        if ' - ' in line_without_year or '–' in line_without_year:
            sep = ' – ' if '–' in line_without_year else ' - '
            parts = [x.strip() for x in line_without_year.split(sep, 1)]
            if len(parts) == 2:
                # Could be "Degree - School" or "School - Degree"
                # Try to detect which is which
                if any(deg_term in parts[0].lower() for deg_term in ['bachelor', 'master', 'phd', 'diploma', 'certificate', 'associate', 'degree', 'b.', 'm.']):
                    degree = parts[0]
                    school = parts[1]
                elif any(deg_term in parts[1].lower() for deg_term in ['bachelor', 'master', 'phd', 'diploma', 'certificate', 'associate', 'degree', 'b.', 'm.']):
                    school = parts[0]
                    degree = parts[1]
                else:
                    # Assume first is degree, second is school
                    degree = parts[0]
                    school = parts[1]
        elif ' | ' in line_without_year:
            parts = [x.strip() for x in line_without_year.split('|')]
            if len(parts) == 2:
                school = parts[0]
                degree = parts[1]
            elif len(parts) >= 3:
                school = parts[0]
                degree = parts[1]
        else:
            # Can't parse cleanly, treat whole line as degree
            degree = line_without_year
        
        # Only add if we have at least degree or school
        if degree or school:
            education.append({
                'degree': degree or 'Qualification',
                'school': school or 'Institution',
                'year': year or 'Present'
            })
    
    return education


def parse_projects_section(text):
    """Parse projects section into structured format.
    
    Handles formats:
    - Project Name
      - Description
    - Project Name - Description
    - Project Name | Description
    
    Returns list of {name, desc, technologies[]}.
    """
    if not text or not text.strip():
        return []
    
    projects = []
    current_project = None
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            if current_project and current_project.get('name'):
                projects.append(current_project)
                current_project = None
            continue
        
        if line.startswith(('-', '•', '*')):
            # Bullet point - treat as description or tech
            content = line.lstrip('-•* ').strip()
            if current_project:
                if not current_project.get('desc'):
                    current_project['desc'] = content
                elif 'technologies' not in current_project:
                    current_project['technologies'] = [content]
                else:
                    current_project['technologies'].append(content)
            continue
        
        # Non-bullet line - could be project name or name-desc combination
        if ' - ' in line or '–' in line or ' | ' in line:
            sep = ' – ' if '–' in line else (' | ' if ' | ' in line else ' - ')
            parts = [x.strip() for x in line.split(sep, 1)]
            
            if current_project and current_project.get('name'):
                projects.append(current_project)
            
            current_project = {
                'name': parts[0],
                'desc': parts[1] if len(parts) > 1 else '',
                'technologies': []
            }
        else:
            # Simple project name
            if current_project and current_project.get('name'):
                projects.append(current_project)
            
            current_project = {
                'name': line,
                'desc': '',
                'technologies': []
            }
    
    # Add last project
    if current_project and current_project.get('name'):
        projects.append(current_project)
    
    # Clean up empty entries
    return [p for p in projects if p.get('name')]


def _generate_suggestions(structured: Dict[str, object], missing_keywords: Sequence[str]) -> List[Dict[str, str]]:
    """Create actionable suggestions based on structured content gaps."""
    suggestions: List[Dict[str, str]] = []

    summary = structured.get("professional_summary") or ""
    if len(summary.split()) < 25:
        suggestions.append({
            "category": "summary",
            "message": "Expand the professional summary to highlight 2-3 quantifiable achievements and core strengths.",
        })

    skills = structured.get("skills", {}) if isinstance(structured.get("skills"), dict) else {}
    if not skills.get("technical"):
        suggestions.append({
            "category": "skills",
            "message": "Add a dedicated technical skills line with relevant tools, frameworks, and platforms.",
        })
    if not skills.get("soft"):
        suggestions.append({
            "category": "skills",
            "message": "Include soft skills such as leadership, stakeholder communication, or team collaboration.",
        })

    experience = structured.get("work_experience", []) or []
    if not experience:
        suggestions.append({
            "category": "experience",
            "message": "Add recent work experience with job titles, companies, dates, and impact-driven bullet points.",
        })
    else:
        for job in experience:
            points = job.get("points") or []
            if len(points) < 2:
                suggestions.append({
                    "category": "experience",
                    "message": f"Add at least two bullet points quantifying impact for {job.get('title') or 'your roles'}.",
                })
                break

    if structured.get("projects") and all(not (proj.get("desc") or proj.get("technologies")) for proj in structured.get("projects", [])):
        suggestions.append({
            "category": "projects",
            "message": "Provide concise descriptions for projects, emphasizing scope, tech stack, and outcomes.",
        })

    if structured.get("education"):
        missing_dates = any(not edu.get("year") for edu in structured.get("education"))
        if missing_dates:
            suggestions.append({
                "category": "education",
                "message": "Add graduation years or expected completion dates for each education entry.",
            })
    else:
        suggestions.append({
            "category": "education",
            "message": "List your highest qualifications with institution names and completion years.",
        })

    if missing_keywords:
        suggestions.append({
            "category": "keywords",
            "message": f"Incorporate role-specific keywords such as: {', '.join(missing_keywords[:10])}.",
        })

    if not structured.get("languages"):
        suggestions.append({
            "category": "languages",
            "message": "Add a languages section if you are proficient in multiple languages relevant to the role.",
        })

    additional = structured.get("additional_information") or ""
    if additional and len(additional.split()) > 120:
        suggestions.append({
            "category": "format",
            "message": "Condense additional information into short bullet points to maintain readability.",
        })

    return suggestions


def convert_to_template_format(sections):
    """Convert raw sections dict into format expected by ResumeTemplate.jsx.
    
    Improved extraction of contact info and structured parsing of all sections.
    """
    # Defensive: ensure sections is a dict (AI may return a string or null for sections)
    if not isinstance(sections, dict):
        sections = {}
    
    # Extract contact info from about/summary section
    about = sections.get('about', '').strip()
    
    # Initialize contact info
    contact_info = {
        'name': 'Your Name',
        'email': 'email@example.com',
        'phone': '+1 (555) 000-0000',
        'location': 'City, Country',
        'address': '',
        'dob': ''
    }
    
    # Try to extract contact details from first lines
    if about:
        lines = about.split('\n')
        extracted = extract_contact_info(about)
        contact_info['name'] = extracted['name'] or contact_info['name']
        contact_info['email'] = extracted['email'] or contact_info['email']
        contact_info['phone'] = extracted['phone'] or contact_info['phone']
        contact_info['location'] = extracted['location'] or contact_info['location']
        contact_info['address'] = extracted.get('address') or contact_info['address'] or extracted.get('location') or contact_info['location']
        contact_info['dob'] = extracted.get('date_of_birth') or contact_info['dob']
        
        # Remove contact info from summary text (keep only the actual summary)
        summary_lines = []
        for line in lines:
            # Skip email, phone, location lines
            if not any(re.search(pattern, line) for pattern in [
                r'[\w.+-]+@[\w-]+\.[\w.-]+',
                r'\+?[\d\s\-()]{10,}',
                r'^(location|city|address|based in):', 
                r'^(phone|mobile|email|website):'
            ]):
                if DOB_LABEL_PATTERN.search(line.lower()):
                    continue
                summary_lines.append(line)
        
        about = '\n'.join(summary_lines).strip()
    
    # Build labeled contact string for readability
    contact_parts = []
    if contact_info.get('email'):
        contact_parts.append(f"email: {contact_info['email']}")
    if contact_info.get('phone') and contact_info['phone'] != '+1 (555) 000-0000':
        contact_parts.append(f"phone: {contact_info['phone']}")
    if contact_info.get('dob'):
        contact_parts.append(f"dob: {contact_info['dob']}")
    if contact_info.get('location') and contact_info['location'] != 'City, Country':
        contact_parts.append(f"location: {contact_info['location']}")
    if contact_info.get('address') and contact_info['address'] not in (contact_info.get('location'), ''):
        contact_parts.append(f"address: {contact_info['address']}")
    contact_str = " | ".join(contact_parts)
    
    # Parse skills into list
    skills_text = sections.get('skills', '').strip()
    skills = [s.strip() for s in re.split(r'[,;]|\n', skills_text) if s.strip()]
    
    # Parse structured sections
    template_data = {
        'name': contact_info['name'] or '',
        'contact': contact_str or '',
        'summary': about or '',
        'skills': skills or [],
        'experience': parse_experience_section(sections.get('experience', '')),
        'projects': parse_projects_section(sections.get('projects', '')),
        'education': parse_education_section(sections.get('education', '')),
        'certifications': []  # Optional section
    }

    # Add certifications from dedicated or achievements sections
    cert_sources = (
        sections.get('certifications', ''),
        sections.get('achievements', ''),
    )
    certs: List[str] = []
    for source in cert_sources:
        if not source:
            continue
        for line in source.split('\n'):
            entry = line.strip().lstrip('-•* ')
            if entry:
                certs.append(entry)
    if certs:
        template_data['certifications'] = _dedupe_preserve_order(certs)

    return template_data


def build_extracted_sections(cv_text: str, structured_sections: Optional[Dict[str, object]] = None) -> Dict[str, object]:
    """Return a cleaned, structured extracted representation of the CV."""
    structured = structured_sections or build_standardized_sections(cv_text or "")

    contact_info = structured.get("contact_information", {}) if isinstance(structured.get("contact_information"), dict) else {}
    skills_dict = structured.get("skills", {}) if isinstance(structured.get("skills"), dict) else {}

    extracted = {
        "header": {
            "name": contact_info.get("name") or "",
            "email": contact_info.get("email") or "",
            "phone": contact_info.get("phone") or "",
            "location": contact_info.get("location") or "",
            "address": contact_info.get("address") or contact_info.get("location") or "",
            "date_of_birth": contact_info.get("dob") or contact_info.get("date_of_birth") or "",
            "linkedin": contact_info.get("linkedin") or "",
            "github": contact_info.get("github") or "",
            "website": contact_info.get("website") or "",
        },
        "professional_summary": structured.get("professional_summary") or "",
        "skills": _dedupe_preserve_order(
            (skills_dict.get("technical") or []) + (skills_dict.get("soft") or []) + (skills_dict.get("other") or [])
        ),
        "experience": structured.get("work_experience") or [],
        "projects": structured.get("projects") or [],
        "education": structured.get("education") or [],
        "certifications": structured.get("certifications") or [],
        "achievements": structured.get("achievements") or [],
        "languages": structured.get("languages") or [],
        "volunteer": structured.get("volunteer_experience") or [],
        "additional": {
            "other": structured.get("additional_information") or "",
        },
    }

    return extracted


def extract_contact_info(text):
    """Extract contact information using NLP and specialized libraries."""
    contact = {
        'name': '',
        'email': '',
        'phone': '',
        'location': '',
        'address': '',
        'date_of_birth': '',
        'linkedin': '',
        'github': '',
        'website': ''
    }
    
    if not text:
        return contact

    # 1. Extract Email (Regex is best)
    email_pattern = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
    email_match = email_pattern.search(text)
    if email_match:
        contact['email'] = email_match.group(0)

    # 2. Extract Phone (phonenumbers lib) with multiple fallbacks
    try:
        for match in phonenumbers.PhoneNumberMatcher(text, "US"):  # Default region US, but finds international too
            contact['phone'] = phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
            break  # Take first valid phone
    except Exception:
        # ignore and fall through to regex-based fallbacks
        pass

    # If phonenumbers didn't find anything, try label-based extraction (e.g., 'Phone: ...')
    if not contact['phone']:
        labeled = re.search(r'(?m)^(?:phone|mobile|tel|telephone)\s*[:\-]\s*(?P<num>.+)$', text, flags=re.IGNORECASE)
        if labeled:
            candidate = labeled.group('num').strip()
            # Keep only common phone characters
            phone_clean = re.sub(r"[^0-9+()\- ]+", "", candidate)
            if phone_clean:
                contact['phone'] = phone_clean

    # Generic regex fallback: look for groups with at least 7 digits (allow spaces/()-)
    if not contact['phone']:
        phone_pattern = re.compile(r'\+?[\d\s\-()]{7,}\d')
        match = phone_pattern.search(text)
        if match:
            contact['phone'] = match.group(0).strip()

    # Last-resort: extract any contiguous digit groups of length >=7
    if not contact['phone']:
        digit_group = re.search(r'(\+?\d[\d\-() ]{6,}\d)', text)
        if digit_group:
            contact['phone'] = digit_group.group(0).strip()

    # 3. Extract Links (Regex)
    linkedin_pattern = re.compile(r'(?:https?://|www\.)?linkedin\.com/[\w\-/]+', re.IGNORECASE)
    github_pattern = re.compile(r'(?:https?://|www\.)?github\.com/[\w\-/]+', re.IGNORECASE)
    url_pattern = re.compile(r'(?:https?://|www\.)[\w./-]+', re.IGNORECASE)
    
    li_match = linkedin_pattern.search(text)
    if li_match: contact['linkedin'] = _normalize_url(li_match.group(0))
    
    gh_match = github_pattern.search(text)
    if gh_match: contact['github'] = _normalize_url(gh_match.group(0))
    
    # 4. Extract Name (NLP)
    # Use Spacy to find PERSON entities in the first few lines
    first_lines = "\n".join(text.split('\n')[:10])
    entities = extract_entities(first_lines)
    if entities.get("PERSON"):
        # Heuristic: Name is usually at the top and not a common word
        for name in entities["PERSON"]:
            if len(name.split()) >= 2 and "@" not in name and not any(char.isdigit() for char in name):
                contact['name'] = name
                break
    
    # Fallback for name if NLP fails (use existing heuristic)
    if not contact['name']:
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            candidate = lines[0]
            # Remove common leading labels like 'Name:' or 'Full Name -'
            candidate = re.sub(r'^(name\s*[:\-]\s*)', '', candidate, flags=re.IGNORECASE).strip()
            if len(candidate.split()) <= 6 and not any(char.isdigit() for char in candidate):
                contact['name'] = candidate

    # 5. Extract Location (NLP GPE)
    if entities.get("GPE"):
        # Additional validation: GPE should not be a single tech term
        location_candidate = entities["GPE"][0]
        # Location should be reasonable (not too short, contains letters)
        if len(location_candidate) >= 3 and not location_candidate.lower() in ['react', 'vue', 'node', 'java', 'php', 'ruby', 'go', 'rust']:
            contact['location'] = location_candidate
            contact['address'] = contact['location']
    
    # Fallback: look for Address: label to capture location/address
    if not contact.get('location'):
        addr_match = re.search(r'(?m)^address\s*[:\-]\s*(?P<addr>.+)$', text, flags=re.IGNORECASE)
        if addr_match:
            addr = addr_match.group('addr').strip()
            contact['address'] = addr
            # try to extract city part (first two comma-separated parts)
            parts = [p.strip() for p in addr.split(',') if p.strip()]
            if parts:
                contact['location'] = parts[0] if len(parts) == 1 else parts[-2] if len(parts) >= 2 else parts[0]

    # 6. Extract DOB if present
    dob = _extract_dob_from_text(text)
    if dob:
        contact['date_of_birth'] = dob

    return contact
    

def _clean_text(value: Optional[str]) -> str:
    if not value:
        return ""
    stripped = value.strip()
    return "" if stripped.lower() == "not provided" else stripped


def _clean_list(values: Optional[Sequence[str]]) -> List[str]:
    cleaned: List[str] = []
    if not values:
        return cleaned
    for value in values:
        if value is None:
            continue
        text = _clean_text(str(value))
        if text:
            cleaned.append(text)
    return cleaned


def _normalize_url(url: Optional[str]) -> str:
    if not url:
        return ""
    cleaned = url.strip().strip('.,);')
    if cleaned and not cleaned.lower().startswith(("http://", "https://")):
        cleaned = "https://" + cleaned.lstrip('/')
    return cleaned


def _clean_experience_entries(entries: Optional[Sequence[Dict[str, object]]]) -> List[Dict[str, object]]:
    cleaned_entries: List[Dict[str, object]] = []
    if not entries:
        return cleaned_entries
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        cleaned_entry = {
            "title": _clean_text(entry.get("title") if isinstance(entry.get("title"), str) else str(entry.get("title") or "")),
            "company": _clean_text(entry.get("company") if isinstance(entry.get("company"), str) else str(entry.get("company") or "")),
            "dates": _clean_text(entry.get("dates") if isinstance(entry.get("dates"), str) else str(entry.get("dates") or "")),
            "location": _clean_text(entry.get("location") if isinstance(entry.get("location"), str) else str(entry.get("location") or "")),
            "points": _clean_list(entry.get("points")),
        }
        # Preserve organization field if available
        if entry.get("organization"):
            cleaned_entry["organization"] = _clean_text(
                entry.get("organization") if isinstance(entry.get("organization"), str) else str(entry.get("organization") or "")
            )
        if any(cleaned_entry.values()) or cleaned_entry.get("points"):
            cleaned_entries.append(cleaned_entry)
    return cleaned_entries


def _clean_project_entries(entries: Optional[Sequence[Dict[str, object]]]) -> List[Dict[str, object]]:
    cleaned_projects: List[Dict[str, object]] = []
    if not entries:
        return cleaned_projects
    for project in entries:
        if not isinstance(project, dict):
            continue
        cleaned_project = {
            "name": _clean_text(project.get("name") if isinstance(project.get("name"), str) else str(project.get("name") or "")),
            "description": _clean_text(
                project.get("desc") if isinstance(project.get("desc"), str) else project.get("description") if isinstance(project.get("description"), str) else str(project.get("desc") or project.get("description") or "")
            ),
            "technologies": _clean_list(project.get("technologies") or project.get("tech")),
        }
        if any((cleaned_project["name"], cleaned_project["description"], cleaned_project["technologies"])):
            cleaned_projects.append(cleaned_project)
    return cleaned_projects


def _clean_education_entries(entries: Optional[Sequence[Dict[str, object]]]) -> List[Dict[str, object]]:
    cleaned_education: List[Dict[str, object]] = []
    if not entries:
        return cleaned_education
    for edu in entries:
        if not isinstance(edu, dict):
            continue
        cleaned_entry = {
            "degree": _clean_text(edu.get("degree") if isinstance(edu.get("degree"), str) else str(edu.get("degree") or "")),
            "school": _clean_text(
                edu.get("school") if isinstance(edu.get("school"), str) else edu.get("institution") if isinstance(edu.get("institution"), str) else str(edu.get("school") or edu.get("institution") or "")
            ),
            "year": _clean_text(edu.get("year") if isinstance(edu.get("year"), str) else edu.get("date") if isinstance(edu.get("date"), str) else str(edu.get("year") or edu.get("date") or "")),
        }
        if any(cleaned_entry.values()):
            cleaned_education.append(cleaned_entry)
    return cleaned_education


def build_structured_cv_payload(structured: Dict[str, object]) -> Dict[str, object]:
    """Build the JSON-ready structured CV output aligned with mandatory sections."""
    structured = structured or {}

    contact = structured.get("contact_information") if isinstance(structured.get("contact_information"), dict) else {}
    logger.info(f"🔍 Building structured payload - contact_information: {contact}")
    logger.info(f"🔍 Raw name value: '{contact.get('name')}', Raw email: '{contact.get('email')}'")
    
    contact_payload = {
        "name": _clean_text(contact.get("name")),
        "email": _clean_text(contact.get("email")),
        "phone": _clean_text(contact.get("phone")),
        "location": _clean_text(contact.get("location")),
        "address": _clean_text(contact.get("address")),
        "date_of_birth": _clean_text(contact.get("dob") or contact.get("date_of_birth")),
        "linkedin": _clean_text(contact.get("linkedin")),
        "github": _clean_text(contact.get("github")),
        "website": _clean_text(contact.get("website")),
    }

    skills_section = structured.get("skills") if isinstance(structured.get("skills"), dict) else {}
    technical_skills = _clean_list(skills_section.get("technical"))
    soft_skills = _clean_list(skills_section.get("soft"))
    other_skills = _clean_list(skills_section.get("other"))
    formatted_skills = _clean_text(skills_section.get("formatted"))
    all_skills = _dedupe_preserve_order(technical_skills + soft_skills + other_skills)

    structured_skills = {
        "technical": technical_skills,
        "soft": soft_skills,
        "other": other_skills,
        "formatted": formatted_skills,
        "all": all_skills,
    }

    structured_payload = {
        "Personal Information": contact_payload,
        "Summary / Objective": _clean_text(structured.get("professional_summary")),
        "Skills": structured_skills,
        "Work Experience / Employment History": _clean_experience_entries(structured.get("work_experience")),
        "Projects": _clean_project_entries(structured.get("projects")),
        "Education": _clean_education_entries(structured.get("education")),
        "Certifications": _clean_list(structured.get("certifications")),
        "Achievements / Awards": _clean_list(structured.get("achievements")),
        "Languages": _clean_list(structured.get("languages")),
        "Volunteer Experience": _clean_experience_entries(structured.get("volunteer_experience")),
        "Additional Information": _clean_text(structured.get("additional_information")),
    }

    # Ensure placeholders are present if sections empty
    for key, value in structured_payload.items():
        if isinstance(value, str) and not value:
            structured_payload[key] = ""
        elif isinstance(value, list) and not value:
            structured_payload[key] = []
        elif isinstance(value, dict) and not any(v for v in value.values() if v):
            structured_payload[key] = {k: v for k, v in value.items()}

    return structured_payload
def optimize_cv_with_gemini(cv_text, job_domain=None):
    """Generate ATS-optimized CV using Gemini AI with proper keywords and formatting.

    This is called when ATS score < 75. Generates a professionally formatted CV
    with domain-specific keywords, action verbs, and proper structure.
    """
    model = get_generative_model()
    if not model:
        raise RuntimeError("AI not configured")

    # Get domain-specific keywords to inject
    domain_keywords = DOMAIN_KEYWORDS.get(job_domain.lower().replace(" ", "_"), []) if job_domain else []
    keywords_hint = f"\n\nIMPORTANT: Integrate these domain-specific keywords naturally: {', '.join(domain_keywords[:15])}" if domain_keywords else ""

    prompt = f"""
    Act as an expert ATS CV optimizer and professional resume writer.
    
    Your task: Transform this CV into a TOP-TIER, ATS-optimized resume that will score 85+ on ATS systems.
    
    Target Domain/Role: {job_domain or 'General Professional'}
    {keywords_hint}

    CRITICAL REQUIREMENTS:

    1. **Content Quality**
       - Remove ALL duplicates, broken formatting, unclear text
       - Fix grammar, punctuation, and sentence structure
       - Use STRONG action verbs: Led, Developed, Implemented, Managed, Optimized, Increased, Reduced, Built, Designed, Launched
       - Add quantifiable achievements with numbers/percentages wherever possible
       - Make every bullet point impactful and results-oriented

    2. **ATS-Friendly Structure** (EXACT ORDER):
       ```
       [FULL NAME]
       Email: email@example.com | Phone: +1-XXX-XXX-XXXX | Location: City, Country
       LinkedIn: url | GitHub: url (if available)
       
       PROFESSIONAL SUMMARY
       3-4 powerful sentences highlighting key strengths, years of experience, and core expertise.
       Must include relevant keywords for the target domain.
       
       TECHNICAL SKILLS
       • Programming Languages: [list]
       • Frameworks & Libraries: [list]
       • Tools & Platforms: [list]
       • Databases: [list]
       • Cloud & DevOps: [list]
       
       PROFESSIONAL EXPERIENCE
       [Job Title] | [Company Name] | [Start Date] - [End Date]
       • [Achievement with metric/number]
       • [Achievement with action verb]
       • [Technical implementation detail]
       (3-5 bullets per role)
       
       PROJECTS (if applicable)
       [Project Name] | [Tech Stack]
       • [What you built and impact]
       • [Measurable result or scale]
       
       EDUCATION
       [Degree] in [Field] — [University Name], [Year]
       
       CERTIFICATIONS (if applicable)
       • [Certification Name] — [Issuer], [Year]
       
       ACHIEVEMENTS & AWARDS (if applicable)
       • [Achievement description]
       ```

    3. **Formatting Rules**
       - Use ONLY plain text formatting (no markdown)
       - Clear section headers (ALL CAPS or Title Case)
       - Consistent bullet points (•, -, or *)
       - NO tables, columns, graphics, or special characters
       - Proper spacing between sections
       - 1-2 pages maximum (400-800 words)

    4. **Keyword Optimization**
       - Naturally integrate domain-specific keywords throughout
       - Include industry-standard terms and technologies
       - Mirror job posting language where applicable
       - Ensure keywords appear in context, not stuffed

    5. **Quality Checklist**
       ✓ Contact information complete
       ✓ Professional summary with keywords
       ✓ Skills categorized clearly
       ✓ Experience with dates and metrics
       ✓ Action verbs in every bullet
       ✓ Quantifiable achievements
       ✓ Education included
       ✓ No spelling/grammar errors
       ✓ ATS-friendly format

    ===============================
    ### JSON OUTPUT
    ===============================
    Return ONLY this JSON structure:
    {{
      "optimized_text": "The complete optimized CV text as a single string following the structure above",
      "sections": {{
        "contact": "Name and contact details",
        "summary": "Professional summary text",
        "skills": "Skills section text",
        "experience": "Work experience section text",
        "education": "Education section text",
        "projects": "Projects section text (if any)"
      }},
      "suggestions": [
        {{"category": "Improvement", "message": "What was improved"}},
        {{"category": "Keywords", "message": "Keywords added: X, Y, Z"}},
        {{"category": "Achievements", "message": "Quantified achievements in role X"}}
      ],
      "keywords_added": ["keyword1", "keyword2", "keyword3"],
      "action_verbs_used": ["Led", "Developed", "Implemented"],
      "estimated_ats_score": 85
    }}

    INPUT CV TO OPTIMIZE:
    {cv_text}
    """

    response = generate_with_retry(model, prompt, generation_config={"response_mime_type": "application/json"})
    if not response:
        raise RuntimeError("AI generation failed after retries")

    try:
        result = json.loads(response.text)
        # Ensure we have the required fields
        if not result.get("optimized_text"):
            raise ValueError("Missing optimized_text in response")
        return result
    except Exception as e:
        logger.error("Failed to parse Gemini JSON response: %s", e)
        # Fallback to manual extraction if JSON mode failed
        content = response.text
        start = content.find('{')
        end = content.rfind('}')
        if start != -1 and end != -1 and end > start:
            return json.loads(content[start:end+1])
        raise


def optimize_cv_with_openai(cv_text: str, job_domain: Optional[str] = None) -> Dict[str, object]:
    """Placeholder - AI parser removed. Use Gemini AI instead."""
    logger.warning("AI CV Parser removed. Use Gemini-based optimization instead.")
    return {}


def optimize_cv(cv_text, job_domain=None, use_ai=True):
    """Unified optimizer: try AI (if requested) and fall back to rule-based optimizer.

    Priority:
    1. OpenAI GPT-4 (best quality, requires API key)
    2. Google Gemini (good quality, fallback)
    3. Rule-based (always works, basic quality)

    Returns dict with keys: optimized_text, sections, template_data, suggestions,
    ats_score, recommended_keywords, found_keywords
    """
    logger.info(f"Starting CV optimization (use_ai={use_ai}, job_domain={job_domain})")
    logger.info(f"CV text length: {len(cv_text) if cv_text else 0} characters")
    
    # Log CV content preview for debugging
    if cv_text:
        preview = cv_text[:500].replace('\n', ' ')
        logger.debug(f"CV content preview: {preview}...")
    
    ai_data: Dict[str, object] = {}
    if use_ai:
        # Use Gemini AI for optimization
        try:
            logger.info("Attempting CV optimization with Gemini AI...")
            # Use Gemini-based optimization here
            ai_data = {}
        except Exception as exc:
            logger.warning(f"✗ Gemini optimization failed: {exc}")
            ai_data = {}
    
    # If AI didn't work or wasn't requested, fall back to rule-based
    rule_based = optimize_cv_rule_based(cv_text, job_domain)

    # Merge AI insights (if any)
    merged = {**rule_based}
    if ai_data:
        # Prefer AI-generated content for optimization if available
        if ai_data.get("optimized_text"):
            merged["optimized_text"] = ai_data["optimized_text"]
            merged["optimized_cv"] = ai_data["optimized_text"]
            merged["optimized_ats_cv"] = ai_data["optimized_text"]

        # Update sections if AI provided structured content
        if ai_data.get("sections") and isinstance(ai_data["sections"], dict):
            merged["sections"] = ai_data["sections"]
            
            # Rebuild ordered_sections based on AI sections to ensure UI consistency
            ai_ordered = []
            
            # Map standard keys to potential AI keys
            key_mapping = {
                "contact_information": ["contact_information", "personal_info", "contact"],
                "professional_summary": ["professional_summary", "summary", "objective", "profile"],
                "work_experience": ["work_experience", "experience", "employment_history"],
                "education": ["education", "academics"],
                "skills": ["skills", "core_competencies", "technical_skills"],
                "projects": ["projects", "key_projects"],
                "certifications": ["certifications", "credentials"],
                "achievements": ["achievements", "awards", "accomplishments"],
                "languages": ["languages"],
                "volunteer_experience": ["volunteer_experience", "volunteering"],
                "additional_information": ["additional_information", "other"]
            }

            # Use standard order for known sections
            for key, label in STANDARD_SECTION_ORDER:
                content = None
                possible_keys = key_mapping.get(key, [key])
                for ai_key in possible_keys:
                    if ai_data["sections"].get(ai_key):
                        content = ai_data["sections"][ai_key]
                        break
                
                if content:
                    if isinstance(content, list):
                        content = "\n".join(str(x) for x in content)
                    ai_ordered.append({"key": key, "label": label, "content": str(content)})
            
            # Add any other sections found in AI response that weren't in standard order
            used_ai_keys = set()
            for v in key_mapping.values():
                used_ai_keys.update(v)
            
            for key, content in ai_data["sections"].items():
                if key not in used_ai_keys:
                    if isinstance(content, list):
                        content = "\n".join(str(x) for x in content)
                    ai_ordered.append({"key": key, "label": key.replace("_", " ").title(), "content": str(content)})
            
            if ai_ordered:
                merged["ordered_sections"] = ai_ordered

        # Update scores and keywords
        for key in ("ats_score", "recommended_keywords", "found_keywords"):
            if key in ai_data:
                merged[key] = ai_data[key]
        if ai_data.get("suggestions"):
            current = [s for s in merged.get("suggestions", []) if isinstance(s, dict)]
            ai_suggestions: List[Dict[str, str]] = []
            for suggestion in ai_data.get("suggestions", []):
                if isinstance(suggestion, dict) and suggestion.get("message"):
                    ai_suggestions.append(suggestion)
                elif isinstance(suggestion, str):
                    ai_suggestions.append({"category": "ai", "message": suggestion})

            encoded = _dedupe_preserve_order(
                [json.dumps(s, sort_keys=True) for s in current] +
                [json.dumps(s, sort_keys=True) for s in ai_suggestions]
            )
            merged["suggestions"] = [json.loads(item) for item in encoded]

    # Validate and log structured payload quality
    if "structured" in merged:
        _validate_structured_payload(merged["structured"])
    
    logger.info("CV optimization complete")
    return merged


def _validate_structured_payload(structured: Dict[str, Any]) -> None:
    """Validate and log quality of structured CV data."""
    if not structured:
        logger.warning("⚠ Structured payload is empty!")
        return
    
    # Check contact information
    contact = structured.get("contact_information", {})
    if isinstance(contact, dict):
        name = contact.get("name", "").strip()
        email = contact.get("email", "").strip()
        phone = contact.get("phone", "").strip()
        
        logger.info(f"Contact Info - Name: {'✓' if name else '✗'}, Email: {'✓' if email else '✗'}, Phone: {'✓' if phone else '✗'}")
        
        if not name:
            logger.warning("⚠ Name not extracted!")
        if not email and not phone:
            logger.warning("⚠ No contact info (email/phone) extracted!")
    else:
        logger.warning("⚠ Contact information is not a dict!")
    
    # Check skills
    skills = structured.get("skills", {})
    if isinstance(skills, dict):
        all_skills = skills.get("all", [])
        technical = skills.get("technical", [])
        logger.info(f"Skills - Total: {len(all_skills)}, Technical: {len(technical)}")
        
        if not all_skills:
            logger.warning("⚠ No skills extracted!")
    else:
        logger.warning("⚠ Skills is not a dict!")
    
    # Check experience
    experience = structured.get("work_experience", [])
    if isinstance(experience, list):
        logger.info(f"Experience - {len(experience)} entries")
        if not experience:
            logger.warning("⚠ No work experience extracted!")
    else:
        logger.warning("⚠ Work experience is not a list!")
    
    # Check education
    education = structured.get("education", [])
    if isinstance(education, list):
        logger.info(f"Education - {len(education)} entries")
        if not education:
            logger.warning("⚠ No education extracted!")
    else:
        logger.warning("⚠ Education is not a list!")
    
    # Calculate completeness score
    critical_fields = 0
    total_critical = 6
    
    if contact.get("name"):
        critical_fields += 1
    if contact.get("email") or contact.get("phone"):
        critical_fields += 1
    if skills.get("all"):
        critical_fields += 1
    if experience:
        critical_fields += 1
    if education:
        critical_fields += 1
    if structured.get("professional_summary"):
        critical_fields += 1
    
    completeness = (critical_fields / total_critical) * 100
    logger.info(f"📊 Extraction completeness: {completeness:.0f}% ({critical_fields}/{total_critical} critical fields)")
    
    if completeness < 50:
        logger.error(f"❌ Extraction quality is poor! Only {completeness:.0f}% complete")
    elif completeness < 80:
        logger.warning(f"⚠ Extraction quality is moderate: {completeness:.0f}% complete")
    else:
        logger.info(f"✓ Extraction quality is good: {completeness:.0f}% complete")

